# -*- coding: utf-8 -*-
"""
Czech DOCX Anonymizer – v6.1
- Načítá jména z JSON knihovny (cz_names.v1.json)
- Opraveno: BANK vs OP, falešné osoby, adresy
Výstupy: <basename>_anon.docx / _map.json / _map.txt
"""

import sys, re, json, unicodedata
from typing import Optional, Set
from pathlib import Path
from collections import defaultdict, OrderedDict
from docx import Document

# =============== Utility ===============
INVISIBLE = '\u00ad\u200b\u200c\u200d\u2060\ufeff'

def clean_invisibles(text: str) -> str:
    if not text: return ''
    text = text.replace('\u00a0', ' ')
    return re.sub('['+re.escape(INVISIBLE)+']', '', text)

def normalize_for_matching(text: str) -> str:
    if not text: return ""
    n = unicodedata.normalize('NFD', text)
    no_diac = ''.join(c for c in n if not unicodedata.combining(c))
    return re.sub(r'[^A-Za-z]', '', no_diac).lower()

def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p

def get_text(p) -> str:
    # KRITICKÁ OPRAVA: Hyperlinky (e-maily, URLs) NEJSOU v p.runs!
    # Musíme použít p.text, který zahrnuje i hyperlinky
    # Fallback na runs je pro případ, kdy p.text nefunguje
    text_from_property = p.text or ''
    if text_from_property:
        return text_from_property
    # Fallback: pokud p.text je prázdný, zkus runs
    return ''.join(r.text or '' for r in p.runs) or ''

def set_text(p, s: str):
    # KRITICKÁ OPRAVA: Pokud paragraph obsahuje hyperlinky, musíme zachovat jejich strukturu
    # Nejjednodušší způsob: smazat všechny runs a hyperlinky a vytvořit nový run
    # (Zachování hyperlinkův by bylo složité, ale nejsou potřeba v anonymizovaném dokumentu)

    # Smaž všechny child elementy (runs, hyperlinky, atd.)
    for child in list(p._element):
        p._element.remove(child)

    # Vytvoř nový run s anonymizovaným textem
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run_elem = OxmlElement('w:r')
    text_elem = OxmlElement('w:t')
    text_elem.text = s
    # Zachovat mezery (preserve space)
    text_elem.set(qn('xml:space'), 'preserve')
    run_elem.append(text_elem)
    p._element.append(run_elem)

def preserve_case(surface: str, tag: str) -> str:
    if surface.isupper(): return tag.upper()
    if surface.istitle(): return tag
    return tag

# =============== Načtení knihovny jmen ===============
def load_names_library(json_path: str = "cz_names.v1.json") -> Set[str]:
    try:
        script_dir = Path(__file__).parent if '__file__' in globals() else Path.cwd()
        json_file = script_dir / json_path

        if not json_file.exists():
            print(f"⚠️  VAROVÁNÍ: {json_path} nenalezen v {script_dir}")
            print(f"⚠️  Kontroluji aktuální složku: {Path.cwd()}")
            # Zkus také aktuální složku
            json_file_cwd = Path.cwd() / json_path
            if json_file_cwd.exists():
                json_file = json_file_cwd
                print(f"✓ Nalezen v aktuální složce")
            else:
                print(f"❌ Soubor {json_path} nebyl nalezen!")
                print(f"   Zkopíruj ho do stejné složky jako skript nebo do aktuální složky.")
                print(f"   Používám prázdnou knihovnu - detekce jmen bude omezená!")
                return set()

        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        names = set()

        # Načteme OBOJÍ - originální jména i normalizovaná
        # Originální jména normalizujeme sami pro konzistenci
        if 'firstnames' in data:
            for name in data['firstnames'].get('M', []):
                names.add(normalize_for_matching(name))
            for name in data['firstnames'].get('F', []):
                names.add(normalize_for_matching(name))

        # Přidáme i předpřipravená normalizovaná jména (fallback)
        if 'firstnames_no_diac' in data:
            names.update(data['firstnames_no_diac'].get('M', []))
            names.update(data['firstnames_no_diac'].get('F', []))

        print(f"✓ Načteno {len(names)} jmen z knihovny")
        return names

    except Exception as e:
        print(f"⚠️  Chyba při načítání: {e}")
        return set()

CZECH_FIRST_NAMES = load_names_library()

# =============== Blacklisty ===============
SURNAME_BLACKLIST = {
    # Právní termíny
    'smlouva','smlouvě','smlouvy','smlouvou','článek','článku','články',
    'datum','číslo','adresa','bydliště','průkaz','občanský','rodné','zákon','sb','kč','čr',
    'ustanovení','příloha','titul','oddíl','bod','pověřený','zástupce','nájem','pronájem',
    'byt','nájemci','nájemce','pronajímatel','pronajímateli',
    'užívat','hlásit','nepřenechávat','elektřina','plyn','sconto','bolton','předat','předání',
    'cena','kauce','záloha','platba','sankce','odpovědnost','poškození','opravy','závady',
    'přepis','přepisem','vyúčtování','paušálně','roční','měsíční',

    # Tabulková a běžná slova (KRITICKÁ OPRAVA: zabránit "Položka Stav" = jméno)
    'stav','stavu','stavem','stavy','stavů','stavech',
    'položka','položky','položku','položek','položkám','položkou','položkami',
    'počet','počtu','počtem','počty','popis','popisu','popisem',
    'celkem','součet','výše','hodnota','hodnoty','množství',
    'období','období','měsíc','měsíce','měsíců','měsíci',
    'splatnost','splatnosti','vyúčtování','doklad','dokladu','faktura','faktury',

    # Značky a produkty
    'jena','dominik','ikea','gorenje','bosch','möbelix',

    # Značky aut
    'škoda','skoda','octavia','fabia','rapid','superb','kodiaq','kamiq','scala','enyaq',
    'volkswagen','audi','seat','bmw','mercedes','toyota','honda','ford','opel','renault',
    'peugeot','citroen','fiat','volvo','mazda','nissan','hyundai','kia',

    # Geografické názvy (s i bez diakritiky)
    'praha','brno','ostrava','plzeň','plzen','liberec','olomouc','budějovice','budejovice',
    'hradec','usti','ústí','pardubice','zlín','zlin','havířov','havirov','kladno','most',
    'opava','frýdek','frydek','karviná','karvina','jihlava','teplice','karlovy','vary',
    'děčín','decin','chomutov','prostějov','prostejov','přerov','prerov','jablonec',
    'ves','město','mesto','obec','vesnice','města','mesta','obce','české','ceske','moravské','moravske',
    'labem','králové','hradec králové',

    # Slova často mylně detekovaná jako příjmení (s i bez diakritiky)
    'bytem','bydliště','bydliste','rodné','rodne','číslo','cislo','císlo','čislo',
    'nový','novy','nová','nova','nové','nove','starý','stary','stará','stara','staré','stare',
    'místo','misto','datum','účtu','uctu','částku','castku','petru',

    # KRITICKÁ OPRAVA: Organizace a firmy (zabránit "Česká Finanční" = jméno)
    'banka','banky','banku','bankám','bankou','bankách','finanční','financni','pojišťovna','pojišťovny','pojištovna',
    'energy','energa','energii','energií','energie','energetický','energeticka',
    'moravia','moravská','moravska','moravské','moravske','českomoravská','ceskomoravska',
    'elektromobilita','elektromobility','elektromobilitě','elektromobilitu',
    'společnost','spolecnost','firma','firmy','firmu','firmou','organizace','organizaci',
    'institut','instituce','instituci','korporace','korporaci','koncern','koncernu',
    'holding','holdingu','group','skupiny','skupina','družstvo','družstva',
    'invest','investment','capital','kapitál','kapitalu','partners','consulting',
    # KRITICKÁ OPRAVA: Zdravotnické a právní organizace (audit smluv 13-16)
    'poliklinika','polikliniky','klinika','kliniky','clinic','cliniec','hospital','nemocnice',
    'notářská','notářské','notářský','notarska','notarske','notarsky','notář','notar',
    'data','dat','processing','protection','gdpr','compliance',
    'london','paris','berlin','vienna','met','mayo','mayo cliniec',
    # KRITICKÁ OPRAVA: Obecné termíny (ne osoby)
    'svoboda','svobody','svobodu','svobodou',  # může být příjmení i termín - v kontextu "Svoboda Notářská" je organizace

    # KRITICKÁ OPRAVA: Role a ne-jména (zabránit "Rodiča Petr", "Učitelka Marie")
    'rodič','rodiče','rodiča','rodičů','rodičům','rodičích','rodičem',
    'učitel','učitelka','učitele','učitelů','učitelům','učitelce','učitelkou',
    'žák','žáci','žáka','žáků','žákům','žákem','student','studenta','studentka','studentkou',
    'matka','matky','matce','matkou','otec','otce','otci','otcem',
    'syn','syna','synovi','synové','dcera','dcery','dceři','dcerou'
}

# KRITICKÁ OPRAVA: Přidat do blacklistu i verze bez diakritiky
# (protože normalize_for_matching() odstraňuje diakritiku)
_blacklist_no_diacritics = set()
for word in SURNAME_BLACKLIST:
    normalized = normalize_for_matching(word)
    if normalized and normalized != word:
        _blacklist_no_diacritics.add(normalized)
SURNAME_BLACKLIST.update(_blacklist_no_diacritics)

ROLE_STOP = {
    'pronajímatel','nájemce','dlužník','věřitel','objednatel','zhotovitel',
    'zaměstnanec','zaměstnavatel','ručitel','spoludlužník','jednatel','svědek',
    'statutární','zástupce','pojistník','pojištěný','odesílatel','příjemce',
    'elektřina','vodné','stočné','topení','internet','služba','služby',

    # Tituly a oslovení
    'pan','paní','pán','slečna','pane','panem',
    'ing','mgr','bc','mudr','judr','phdr','rndr','doc','prof','csc','ph','dr'
}

# =============== Inference nominativu ===============
def _male_genitive_to_nominative(obs: str) -> Optional[str]:
    """Převede pozorovaný tvar (např. genitiv) na nominativ pro mužská jména."""
    lo = obs.lower()
    cands = []

    # Dativ/Vokativ: -u → nominativ (Michalu → Michal, Petru → Petr)
    # DŮLEŽITÉ: Testujeme to PŘED -a, protože Michalu má končit na -u, ne -a
    if lo.endswith('u') and len(obs) > 1:
        cands.append(obs[:-1])

    # Dativ/Lokál: -ovi → nominativ (Petrovi → Petr)
    if lo.endswith('ovi') and len(obs) > 3:
        cands.append(obs[:-3])

    # Instrumentál: -em → nominativ (Petrem → Petr)
    if lo.endswith('em') and len(obs) > 2:
        cands.append(obs[:-2])

    # Speciální případy: -ka → -ek, -la → -el, -ca → -ec
    if lo.endswith('ka') and len(obs) > 2:
        cands.append(obs[:-2] + 'ek')
    if lo.endswith('la') and len(obs) > 2:
        cands.append(obs[:-2] + 'el')
    if lo.endswith('ca') and len(obs) > 2:
        cands.append(obs[:-2] + 'ec')

    # Genitiv/Akuzativ: -a → nominativ (Petra → Petr)
    # Testujeme až po -ka/-la/-ca, abychom správně zachytili speciální případy
    if lo.endswith('a') and len(obs) > 1:
        cands.append(obs[:-1])

    # Vokativ/Lokál: -e → nominativ (u Pavle → Pavel)
    if lo.endswith('e') and len(obs) > 1:
        cands.append(obs[:-1])
        # Možné měkčení zpět: Pavle → Pavel
        if len(obs) > 2 and obs[-2:-1] in 'lc':
            cands.append(obs[:-1] + 'el')

    # Kontrola proti knihovně jmen
    for cand in cands:
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand
    return None

def infer_first_name_nominative(observed: str, surname_observed: str = "") -> Optional[str]:
    """
    Odvozuje nominativ křestního jména z pozorovaného tvaru (může být v jakémkoliv pádu).
    Například: "Petra" → "Petr", "Janě" → "Jana", "Jiřího" → "Jiří"
    """
    if not observed: return None
    obs = observed.strip()
    surname_lower = (surname_observed or "").lower()
    female_like_surname = surname_lower.endswith(('ová', 'á', 'ou', 'é'))

    # Zkus nejdřív přímé matchování
    norm = normalize_for_matching(obs)
    if norm in CZECH_FIRST_NAMES:
        return obs

    # Pokud příjmení nenaznačuje ženu, zkus mužská pravidla
    if not female_like_surname:
        cand = _male_genitive_to_nominative(obs)
        if cand: return cand

    # ========== Ženská jména ==========
    low = obs.lower()

    # Speciální případ: -ice → -ika nebo -a (Verunice → Veronika)
    if low.endswith('ice') and len(obs) > 3:
        cand = obs[:-3] + 'ika'
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand
        cand = obs[:-3] + 'a'
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand

    # Speciální případ: -ře → -ra (Petře → Petra)
    if low.endswith('ře') and len(obs) > 2:
        cand = obs[:-2] + 'ra'
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand

    # Přivlastňovací tvary: -in/-ina/-iny/... → -a
    for suf in ['inou','iným','iných','iné','inu','iny','ina','in']:
        if low.endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)] + 'a'
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand

    # Základní pády: -ou/-u/-y/-e/-ě/-o → -a
    for suf in ['ou','u','y','e','ě','o']:
        if low.endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)] + 'a'
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand

    # ========== Mužská jména (alternativní cesta) ==========

    # Přivlastňovací tvary: -ův/-ova/-ovo/-ových/... → základ
    for suf in ['ových','ovou','ově','ovu','ova','ovo','ův']:
        if low.endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)]
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand

    # Základní pády mužských jmen
    for suf in ['ovi','em','e','u','a']:
        if low.endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)]
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand

    # Speciální případ pro jména na -í (Jiří)
    if low.endswith(('ího','ímu','ím','íh')):
        for suf_len in [3, 3, 2, 2]:
            if len(obs) > suf_len:
                cand = obs[:-suf_len] + 'í'
                if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                    return cand

    return None

def infer_surname_nominative(observed: str) -> str:
    """
    Odvozuje nominativ příjmení z pozorovaného tvaru.
    Například: "Novákovi" → "Novák", "Novákovou" → "Nováková", "Novotného" → "Novotný"
    """
    if not observed: return observed
    obs = observed.strip()
    low = obs.lower()

    # ========== Ženská příjmení typu -ová ==========
    if low.endswith('ovou') and len(obs) > 4:
        return obs[:-4] + 'ová'  # Novákovou → Nováková
    if low.endswith('ové') and len(obs) > 3:
        # Může být gen/dat/lok Novákové, ale nominativ je Nováková
        return obs[:-3] + 'ová'
    if low.endswith('ou') and len(obs) > 2 and not low.endswith('ovou'):
        # Instrumentál: Novákovou → Nováková (ale opatrně)
        # Může být i příjmení typu Malou → Malá
        if low.endswith('ovou'):  # už jsme ošetřili výše
            return obs[:-4] + 'ová'
        return obs[:-2] + 'á'

    # ========== Přídavná jména typu -ský/-cký/-ný ==========
    # Novotného → Novotný, Novotné → Novotná/Novotný
    if low.endswith(('ského','ckého')):
        return obs[:-3] + 'ý'  # Novotského → Novotný
    if low.endswith(('ému','ských','ckých','ským','ckým')):
        # Různé pády
        suffix_map = {'ému': 'ý', 'ských': 'ý', 'ckých': 'cký', 'ským': 'ý', 'ckým': 'cký'}
        for s, repl in suffix_map.items():
            if low.endswith(s):
                return obs[:-len(s)] + repl
    if low.endswith('nou'):
        # Instrumentál ženska forma: Novotnou → Novotná, Suchou → Suchá
        # KRITICKÁ OPRAVA: Vyjmout zvířecí příjmení (Vrbou → Vrba, ne Vrbá)
        base = obs[:-2]  # Odstraň -ou
        if base.lower().endswith(('rb', 'rk', 'lk')):
            # Je to pravděpodobně zvířecí příjmení → pokračuj níže
            pass  # Neskoč do return, pokračuj v kódu
        else:
            return obs[:-3] + 'ná'
    # Obecný test pro příjmení končící na '-é' (přídavná jména)
    # Suché, Novotné, Malé, atd. → Suchá, Novotná, Malá
    if low.endswith('é') and len(obs) > 1:
        # Může být gen/dat/lok od ženské formy
        # NEBO nominativ středního rodu (vzácné u příjmení)
        # Pro příjmení předpokládáme ženský tvar
        return obs[:-1] + 'á'
    if low.endswith(('ým','ém')) and len(obs) > 2:
        # Možná instrumentál/lokál -ým/-ém
        return obs[:-1] + 'ý'

    # ========== Speciální případy pro příjmení typu -ček/-nek/-ek ==========
    m = re.match(r'^(.+)ček(a|ovi|em|u|e|y|ou|ům|ách|ů)?$', obs, flags=re.IGNORECASE)
    if m:
        return m.group(1) + 'ček'

    m2 = re.match(r'^(.+)n[eě]k(a|ovi|em|u|e|y|ou|ům|ách|ů)?$', obs, flags=re.IGNORECASE)
    if m2:
        return m2.group(1) + 'nek'

    # DŮLEŽITÉ: Pouze pro příjmení typu -ek (Hájek, Čábelek), NE pro běžná příjmení+'kem' (Dvořákem)
    # KRITICKÁ OPRAVA: NE pro zvířecí příjmení (Liška, ne Lišek)
    # Kontrola: před 'k' musí být souhláska (ne samohláska)
    if low.endswith(('ka','kovi','kem','ku','ke','ků','kům')) and len(obs) > 3:
        # KRITICKÁ OPRAVA: Vyjmout zvířecí příjmení (Liška, ne Lišek)
        # Pokud to vypadá jako zvířecí příjmení, NEPŘEPISOVAT na -ek
        base_without_suffix = re.sub(r'k(ovi|em|u|e|a|ů|ům)?$', 'k', obs, flags=re.IGNORECASE)
        if not base_without_suffix.lower().endswith(('išk', 'íšk', 'ešk', 'ůbk', 'ubk')):
            # Zjisti, který suffix máme
            for suff in ['kovi', 'kem', 'kům', 'ka', 'ku', 'ke', 'ků']:
                if low.endswith(suff):
                    idx_before_k = -(len(suff) + 1)
                    if len(obs) >= abs(idx_before_k):
                        char_before_k = obs[idx_before_k].lower()
                        # Pouze pokud je před 'k' souhláska (příjmení typu Hájek)
                        if char_before_k not in 'aáeéěiíoóuúůyý':
                            return re.sub(r'k(ovi|em|u|e|a|ů|ům)?$', 'ek', obs, flags=re.IGNORECASE)
                    break

    # ========== Příjmení typu -ec (Němec) ==========
    m3 = re.match(r'^(.+)c(e|i|em|ů|ích|ům|ech|emi|u|y)?$', obs, flags=re.IGNORECASE)
    if m3:
        return m3.group(1) + 'ec'

    # ========== Příjmení na -a (Svoboda) ==========
    if low.endswith('ovi') and len(obs) > 3:
        # Svobodovi → Svoboda
        base = obs[:-3]
        # Ale pokud je to -ovi pro příjmení bez -a, pak → základ
        # Zkusíme přidat -a
        return base + 'a'

    # Ostatní pády pro příjmení na -a
    for suf in ['ou','e','u','y']:
        if low.endswith(suf) and len(obs) > len(suf)+1:
            # Svobodou → Svoboda, Svobodě → Svoboda
            base = obs[:-len(suf)]
            # KRITICKÁ OPRAVA: Kontrola zvířecích příjmení (Vrbou → Vrba, ne Vrbá)
            # Pokud základ končí na souhlásku, může být nominativ buď se "-a" nebo bez
            # Vrba je nominativ (zvířecí příjmení), Svoboda je nominativ (obecné)
            # Zkontroluj, zda base + 'a' je v seznamu nominativních vzorů
            test_with_a = (base + 'a').lower()
            if test_with_a.endswith(('rba','íška','iška','ána','vrána','liška','holuba','jelínka')):
                # Je to pravděpodobně zvířecí příjmení → nominativ je base + 'a'
                return base + 'a'
            candidate = base + 'a'
            # Ale pozor: může to být i příjmení bez -a
            # Pokud původní slovo končí na souhlásku, může to být Novák
            return candidate

    # ========== Obecná mužská příjmení (konsonantní kmeny) ==========
    # Novák, Dvořák, Malý, atd.

    # Dativ/Lokál: -ovi, Instrumentál: -em
    if low.endswith('ovi') and len(obs) > 3:
        base = obs[:-3]  # Novákovi → Novák
        # KRITICKÁ OPRAVA: Kontrola vložného e (Havlovi → Havel, ne Havl)
        if len(base) >= 3:
            last_two = base[-2:].lower()
            if last_two in ('vl', 'dl', 'kl', 'pl', 'sl', 'zl', 'čl', 'šl', 'tl', 'hl', 'bl', 'gl'):
                char_before = base[-3].lower()
                if char_before in 'aáeéěiíoóuúůyý':
                    return base[:-1] + 'e' + base[-1:]  # Havl → Havel

        # KRITICKÁ OPRAVA: Příjmení typu Liška, Holub (zvířecí příjmení)
        # Liškovi → Liška (ne Lišk), Holubovi → Holub
        # Heuristika: Pokud základ končí na typické zvířecí vzory, přidej -a
        if base.lower().endswith(('išk', 'íšk', 'ešk', 'ůbk', 'ubk', 'oub', 'lub', 'rán', 'ván')):
            return base + 'a'
        return base
    if low.endswith('em') and len(obs) > 2:
        base = obs[:-2]  # Novákem → Novák
        # KRITICKÁ OPRAVA: Kontrola vložného e (Havlem → Havel, ne Havl)
        if len(base) >= 3:
            last_two = base[-2:].lower()
            if last_two in ('vl', 'dl', 'kl', 'pl', 'sl', 'zl', 'čl', 'šl', 'tl', 'hl', 'bl', 'gl'):
                char_before = base[-3].lower()
                if char_before in 'aáeéěiíoóuúůyý':
                    return base[:-1] + 'e' + base[-1:]  # Havl → Havel
        return base

    # Genitiv/Akuzativ: -a (ale POUZE pokud to není příjmení na -a v nominativu!)
    # Problém: "Říha" je nominativ, ale končí na -a
    # "Nováka" je genitiv od "Novák"
    # Heuristika: Pokud celé slovo končí na typické vzory, je to nominativ
    if low.endswith('a') and len(obs) > 2:
        # Typické vzory pro příjmení v nominativu na -a:
        # -ha, -la, -ra, -da, -ta, -na, -ka, -cha, -ma, -ba, -pa, -va, -za, -sa
        # Příklady: Říha, Skála, Hora, Svoboda, Kučera, Vrána, Liška
        # KRITICKÁ OPRAVA: Přidány vzory s dlouhými samohláskami (ána, íška, ůbka)
        if low.endswith(('iha','íha','uha','ůha','eha','ěha','oha','aha','ána','yha',
                         'ila','íla','ula','ůla','ela','ěla','ola','ala','ála','yla',
                         'ira','íra','ura','ůra','era','ěra','ora','ara','ára','yra',
                         'ida','ída','uda','ůda','eda','ěda','oda','ada','áda','yda',
                         'ita','íta','uta','ůta','eta','ěta','ota','ata','áta','yta',
                         'ina','ína','una','ůna','ena','ěna','ona','ana','ána','yna',
                         'ika','íka','uka','ůka','eka','ěka','oka','aka','áka','yka','íška','iška','ůbka','ubka','ybka',
                         'ima','íma','uma','ůma','ema','ěma','oma','ama','áma','yma',
                         'iba','íba','uba','ůba','eba','ěba','oba','aba','ába','yba','rba',
                         'ipa','ípa','upa','ůpa','epa','ěpa','opa','apa','ápa','ypa',
                         'iva','íva','uva','ůva','eva','ěva','ova','ava','áva','yva',
                         'iza','íza','uza','ůza','eza','ěza','oza','aza','áza','yza',
                         'isa','ísa','usa','ůsa','esa','ěsa','osa','asa','ása','ysa')):
            # Je to pravděpodobně nominativ
            return obs
        else:
            # KRITICKÁ OPRAVA: Vložné e/ě v příjmeních (Havl → Havel, Vrán → Vrána)
            # Pokud základní tvar (po odebrání -a) končí na dvě souhlásky, může to být vložné e
            base = obs[:-1]
            if len(base) >= 2:
                last_two = base[-2:].lower()
                # Běžné vzory s vložným e: -vl, -dl, -kl, -pl, -sl, -zl, -čl, -šl
                # Běžné vzory s vložným ě: -st, -šť, -čt
                if last_two in ('vl', 'dl', 'kl', 'pl', 'sl', 'zl', 'čl', 'šl', 'tl', 'hl', 'bl', 'gl'):
                    # Vlož 'e': Havl → Havel
                    return base[:-1] + 'e' + base[-1:]
                elif last_two in ('st', 'šť', 'čt', 'zt', 'žď'):
                    # Vlož 'ě': možná Štěpánský, ale to už je ošetřeno výše
                    # Pro jistotu necháme bez změny
                    pass
            # Jinak je to pravděpodobně genitiv → odebrat -a
            return base

    # KRITICKÁ OPRAVA: Pokud příjmení v nominativu končí na souhláskovou skupinu,
    # která by měla mít vložné e, přidej ho (Havl → Havel, Vrb → Vrba)
    # Toto řeší chyby v dokumentu, kde je napsáno "Petr Havl" místo "Petr Havel"
    if len(obs) >= 3:
        last_two = obs[-2:].lower()
        if last_two in ('vl', 'dl', 'kl', 'pl', 'sl', 'zl', 'čl', 'šl', 'tl', 'hl', 'bl', 'gl', 'rb', 'mb'):
            # Kontrola, že před tím je samohláska (aby "Navl" nebylo "Navel")
            if len(obs) >= 3:
                char_before = obs[-3].lower()
                if char_before in 'aáeéěiíoóuúůyý':
                    # Vlož 'e': Havl → Havel, Vrb → Vrba (ne, počkat - Vrb → Vrb)
                    # Vlastně ne, Vrb se deklinuje jako Vrba (nom. Vrb, gen. Vrba)
                    # Zkusme jen pro -vl, -dl, -kl, -pl, -sl, -zl, -čl, -šl, -tl, -hl, -bl, -gl
                    if last_two in ('vl', 'dl', 'kl', 'pl', 'sl', 'zl', 'čl', 'šl', 'tl', 'hl', 'bl', 'gl'):
                        return obs[:-1] + 'e' + obs[-1:]

    return obs

# =============== Varianty pro nahrazování ===============
def variants_for_first(first: str) -> set:
    """
    Generuje všechny pádové varianty křestního jména včetně:
    - Nominativ, Genitiv, Dativ, Akuzativ, Vokativ, Lokál, Instrumentál
    - Přivlastňovací přídavná jména (Petrův, Janin, atd.)
    """
    f = first.strip()
    if not f: return {''}
    V = {f, f.lower(), f.capitalize()}
    low = f.lower()

    # ========== Ženská jména končící na -a ==========
    if low.endswith('a'):
        stem = f[:-1]
        # Základní pády: Gen/Dat/Akuz/Vok/Lok/Instr
        V |= {stem+'y', stem+'e', stem+'ě', stem+'u', stem+'ou', stem+'o'}

        # Přivlastňovací přídavná jména (Janin dům, Petřina kniha)
        V |= {stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných','ino']}

        # Speciální případy pro měkčení (Petra → Petře, Veronka → Verunce)
        if stem.endswith('k'):
            V.add(stem[:-1] + 'c' + 'e')  # Veronka → Verunce
            V.add(stem[:-1] + 'c' + 'i')  # Veronka → Verunce (alt)

        # Speciální měkčení tr → tř (Petra → Petřin)
        if stem.endswith('tr'):
            soft_stem = stem[:-1] + 'ř'
            V |= {soft_stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných','ino']}

        # Speciální měkčení h → z, ch → š, k → c, r → ř
        if stem.endswith('h'):
            soft_stem = stem[:-1] + 'z'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')
        if stem.endswith('ch'):
            soft_stem = stem[:-2] + 'š'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')
        if stem.endswith(('k', 'g')):
            soft_stem = stem[:-1] + 'c'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')
        if stem.endswith('r') and not stem.endswith('tr'):
            soft_stem = stem[:-1] + 'ř'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')

    # ========== Mužská jména ==========
    else:
        # Základní pády
        V |= {f+'a', f+'ovi', f+'e', f+'em', f+'u', f+'om'}

        # Přivlastňovací přídavná jména (Petrův dům, Petrova kniha)
        V |= {f+'ův', f+'ova', f+'ovo', f+'ovu', f+'ovou', f+'ově'}
        V |= {f+'ov'+s for s in ['a','o','y','ě','ým','ých','ou','u','e']}

        # Speciální případy pro zakončení -ek, -el
        if low.endswith('ek'):
            stem_k = f[:-2] + 'k'
            V |= {stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e'}
            V.add(f[:-2] + 'ka')  # Vladimírek → Vladimírka

        if low.endswith('el'):
            stem_l = f[:-2] + 'l'
            V |= {stem_l+'a', stem_l+'ovi', stem_l+'em', stem_l+'u', stem_l+'e'}
            V.add(f[:-2] + 'la')  # Pavel → Pavla

        # Speciální případy pro zakončení -ec
        if low.endswith('ec'):
            stem_c = f[:-2] + 'c'
            V |= {stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'u'}

        # Speciální případ: Jiří → Jiřího, Jiřímu, Jiřím, Jiřího
        if low.endswith('í'):
            stem = f[:-1]
            V |= {stem+'ího', stem+'ímu', stem+'ím', stem+'íh'}

        # Speciální případ: -iš/-aš → měkčení (Lukáš, Tomáš)
        if low.endswith(('áš', 'iš')):
            stem_base = f[:-1]
            V |= {stem_base+'e', stem_base+'i', stem_base+'em', stem_base+'ovi'}

        # Lokál s měkčením (Petr → o Petrovi, ale Pavel → o Pavlovi)
        if not low.endswith(('i', 'í')):
            V |= {f+'ovi', f+'e'}  # "o Petrovi", "u Petra"

    # Přidání verzí bez diakritiky
    V |= {unicodedata.normalize('NFKD', v).encode('ascii','ignore').decode('ascii') for v in list(V)}

    return V

def variants_for_surname(surname: str) -> set:
    """
    Generuje všechny pádové varianty příjmení včetně:
    - Všechny pády jednotného i množného čísla
    - Přivlastňovací přídavná jména (Novákův, Novákova)
    - Speciální případy pro -ová, -ský, -ek, -ec, atd.
    """
    s = surname.strip()
    if not s: return {''}
    out = {s, s.lower(), s.capitalize()}
    low = s.lower()

    # ========== Příjmení typu -ová (ženská) ==========
    if low.endswith('ová'):
        base = s[:-1]  # Novákov
        out |= {
            s,              # Nováková (nom)
            base+'é',       # Novákové (gen/dat/lok)
            base+'ou',      # Novákovou (instr)
            base+'á',       # alternativa (nom)
        }
        # Množné číslo
        base_stem = s[:-3]  # Novák
        out |= {
            base_stem+'ových',  # u Novákových (gen pl)
            base_stem+'ovým',   # Novákovým (dat/instr pl)
            base_stem+'ové',    # Novákové (nom pl)
        }
        return out

    # ========== Příjmení typu -ský/-cký (přídavná jména) ==========
    if low.endswith(('ský','cký')):
        stem = s[:-2]  # Novot
        out |= {
            stem+'ý',       # Novotný (nom m)
            stem+'ého',     # Novotného (gen/akuz)
            stem+'ému',     # Novotnému (dat)
            stem+'ým',      # Novotným (instr)
            stem+'ém',      # Novotném (lok)
            stem+'á',       # Novotná (nom f)
            stem+'é',       # Novotné (gen/dat/lok f)
            stem+'ou',      # Novotnou (instr f)
            stem+'ých',     # Novotných (gen pl)
            stem+'ými',     # Novotnými (instr pl)
            stem+'ým',      # Novotným (dat pl)
        }
        return out

    # ========== Obecná přídavná jména končící na -ý ==========
    if low.endswith('ý'):
        stem = s[:-1]
        out |= {
            stem+'ý', stem+'ého', stem+'ému', stem+'ým', stem+'ém',
            stem+'á', stem+'é', stem+'ou',
            stem+'ých', stem+'ými'
        }
        return out

    # ========== Ženská příjmení na -á (ne -ová) ==========
    if low.endswith('á') and not low.endswith('ová'):
        stem = s[:-1]
        out |= {s, stem+'é', stem+'ou', stem+'á'}
        return out

    # ========== Příjmení typu -ek (Dvořáček, Hájek) ==========
    if low.endswith('ek') and len(s) >= 3:
        stem_k = s[:-2] + 'k'
        out |= {
            s,              # Dvořáček (nom)
            stem_k+'a',     # Dvořáčka (gen)
            stem_k+'ovi',   # Dvořáčkovi (dat)
            stem_k+'em',    # Dvořáčkem (instr)
            stem_k+'u',     # Dvořáčku (vok/akuz)
            stem_k+'e',     # Dvořáčku (lok)
            stem_k+'y',     # alt
            stem_k+'ou',    # alt
        }
        # Přivlastňovací
        out |= {
            stem_k+'ův', stem_k+'ova', stem_k+'ovo',
            stem_k+'ovu', stem_k+'ovou', stem_k+'ově'
        }
        # Množné číslo
        out |= {
            stem_k+'ů',     # u Dvořáčků (gen pl)
            stem_k+'ům',    # Dvořáčkům (dat pl)
            stem_k+'y',     # Dvořáčky (akuz pl)
        }
        return out

    # ========== Příjmení typu -ec (Němec, Konec) ==========
    if low.endswith('ec') and len(s) >= 3:
        stem_c = s[:-2] + 'c'
        out |= {
            s,              # Němec (nom)
            stem_c+'e',     # Němce (gen/akuz)
            stem_c+'i',     # Němci (dat/lok)
            stem_c+'em',    # Němcem (instr)
            stem_c+'u',     # alt
            stem_c+'y',     # alt
        }
        # Množné číslo
        out |= {
            stem_c+'ů',     # Němců (gen pl)
            stem_c+'ům',    # Němcům (dat pl)
            stem_c+'ích',   # Němcích (lok pl)
            stem_c+'ech',   # Němcech (alt lok pl)
            stem_c+'emi',   # Němcemi (instr pl)
        }
        # Přivlastňovací
        out |= {
            stem_c+'ův', stem_c+'ova', stem_c+'ovo',
            stem_c+'ovu', stem_c+'ovou', stem_c+'ově'
        }
        return out

    # ========== Příjmení na -a (mužská i ženská) ==========
    if low.endswith('a') and len(s) >= 2 and not low.endswith('ová'):
        stem = s[:-1]
        out |= {
            s,              # Svoboda (nom)
            stem+'y',       # Svobody (gen)
            stem+'ovi',     # Svobodovi (dat m)
            stem+'ou',      # Svobodou (instr)
            stem+'u',       # Svobodu (akuz)
            stem+'e',       # Svobodě (lok)
            stem+'o',       # vok
        }
        # Přivlastňovací
        out |= {
            stem+'ův', stem+'ova', stem+'ovo',
            stem+'ovu', stem+'ovou', stem+'ově'
        }
        # Množné číslo
        out |= {
            stem+'ů',       # u Svobodů (gen pl)
            stem+'ům',      # Svobodům (dat pl)
            stem+'y',       # Svobody (akuz pl)
        }
        return out

    # ========== Obecná mužská příjmení (konsonantní kmeny) ==========
    # Novák, Dvořák, Malý, atd.
    out |= {
        s+'a',          # Nováka (gen)
        s+'ovi',        # Novákovi (dat)
        s+'e',          # Nováku (lok/vok)
        s+'em',         # Novákem (instr)
        s+'u',          # Nováku (alt)
    }
    # Přivlastňovací přídavná jména
    out |= {
        s+'ův', s+'ova', s+'ovo',
        s+'ovu', s+'ovou', s+'ově'
    }
    out |= {
        s+'ov'+suf for suf in ['a','o','y','ě','ým','ých','ou','u','e','i']
    }
    # Množné číslo
    out |= {
        s+'ů',          # u Nováků (gen pl)
        s+'ům',         # Novákům (dat pl)
        s+'y',          # Nováky (akuz pl)
        s+'ích',        # Novácích (lok pl)
        s+'ech',        # alt lok
    }

    # Přidání verzí bez diakritiky
    out |= {unicodedata.normalize('NFKD', v).encode('ascii','ignore').decode('ascii') for v in list(out)}

    return out

# =============== Regexy ===============
# Vylepšený ADDRESS_RE - zachytává čistou adresu (Ulice číslo, PSČ Město)
# Podporuje prefixy: "Sídlo:", "Bytem:", "v ulici", "Místo podnikání:", atd.
# DŮLEŽITÉ: Adresa MUSÍ mít formát "Ulice číslo, Město" (čárka + město jsou povinné)
# VYLUČUJE: formát "Jméno Příjmení, bytem..." (to je osoba + adresa, ne jen adresa)
# KRITICKÁ OPRAVA: Podpora pro zkratky ulic (nám., ul., tř.)
# KRITICKÁ OPRAVA: Prefix je nyní volitelný (např. "IČO: 123456, Na Příkopě 33, Praha 1")
ADDRESS_RE = re.compile(
    r'(?<!\[)'                                       # Ne po '['
    r'(?:'                                           # Začátek prefixů (VOLITELNÉ!)
    r'(?:(?:trvale\s+)?bytem\s*:?\s*)|'             # "bytem" nebo "Bytem:"
    r'(?:(?:trvalé\s+)?bydlišt[eě]\s*:\s*)|'        # "trvalé bydliště:"
    r'(?:(?:sídlo(?:\s+podnikání)?|se\s+sídlem)\s*:\s*)|'  # "sídlo:" / "se sídlem:"
    r'(?:místo\s+(?:podnikání|výkonu\s+práce)\s*:?\s*)|'  # "Místo podnikání:" nebo "Místo výkonu práce" (volitelná :)
    r'(?:(?:adresa|trvalý\s+pobyt)\s*:\s*)|'       # "adresa:" / "trvalý pobyt:"
    r'(?:(?:v\s+ulic[ií]|na\s+adrese|v\s+dom[eě])\s+)'  # "v ulici " / "na adrese " (BEZ volitelnosti!)
    r')?'                                            # CELÝ PREFIX JE VOLITELNÝ!
    r'(?![A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+\s+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+,\s+bytem)'  # VYLUČUJE: "Jméno Příjmení, bytem"
    r'(?![A-Z]{2,3}\s+\d{6,9})'                      # VYLUČUJE: "AB 456789" (OP kódy)
    r'(?:'                                           # Začátek ulice
    r'(?:nám\.|ul\.|tř\.|n\.|u\.|t\.)\s+|'          # Zkratky: nám. (náměstí), ul. (ulice), tř. (třída)
    r'(?:(?:Na|U|K|Pod|V|Nad|Za)\s+)?'              # Volitelné předložky (Na Příkopě, U Lávky, K Lesu, Pod Skalkou)
    r'(?:[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ])'                     # Velké písmeno (začátek názvu)
    r')'
    r'[a-záčďéěíňóřšťúůýž\s]{1,50}?'                # Název ulice (non-greedy OK, ukončeno číslem)
    r'\s+\d{1,4}(?:/\d{1,4})?'                      # Číslo domu (25 nebo 25/8)
    r',\s*'                                          # Čárka POVINNÁ
    r'(?:\d{3}\s?\d{2}\s+)?'                         # PSČ volitelné (612 00)
    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'                         # Velké písmeno (začátek města)
    r'(?:(?:(?!Tel\.?|Nar\.?|Rodn[éě]|Číslo|IČO|DIČ)[a-záčďéěíňóřšťúůýž\s\d\-])+?)'  # Město - negative lookahead, přidána pomlčka pro "Brno-střed"
    r'(?=\s*(?:$|[,.\n()\[\]]|(?:Nar\.?|RČ|Rodn[éě]|IČO|DIČ|OP|Občansk|Tel\.?|Telefon|E-mail|Kontakt|Číslo|Datum|Zastoupen|Jednatel|vyd[aá]n|dále)))',  # Lookahead
    re.UNICODE | re.IGNORECASE
)

# ADDRESS_WITH_ZIP_RE - adresy s PSČ BEZ prefixu (pro tabulky, kde prefix je v jiném cell)
# Formát: "Ulice číslo, PSČ Město" - PSČ je POVINNÉ pro jednoznačnost
# Příklad: "Čechova 14, 750 02 Přerov" v tabulce pod hlavičkou "Adresa trvalého pobytu"
# KRITICKÁ OPRAVA: Vyloučit prefixní fráze jako "NP domu na adrese" z matche - prefix je v group 1, adresa v group 2
ADDRESS_WITH_ZIP_RE = re.compile(
    r'(?<!\[)'                                       # Ne po '['
    r'((?<!\w)(?:v\s+)?(?:\d+\.)?\s*NP\s+(?:domu\s+)?(?:na\s+adrese|v\s+dom[eě]|v\s+ulic[ií])\s+)?'  # Group 1: Volitelný prefix
    r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'                        # Group 2 ZAČÁTEK - Velké písmeno (začátek ulice)
    r'(?![Nn][Pp]\s)'                                # NESMÍ začínat s "NP " nebo "Np "
    r'[a-záčďéěíňóřšťúůýž\s]{2,50}?'                # Název ulice (non-greedy OK, ukončeno číslem)
    r'\s+\d{1,4}(?:/\d{1,4})?'                      # Číslo domu
    r',\s*'                                          # Čárka
    r'\d{3}\s?\d{2}\s+'                              # PSČ POVINNÉ (612 00 nebo 61200)
    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'                         # Velké písmeno (začátek města)
    r'(?:(?:(?!Tel\.?|Nar\.?|Rodn[éě]|Číslo)[a-záčďéěíňóřšťúůýž\s\d])+?)'  # Město - negative lookahead pro klíčová slova
    r')(?=\s*(?:$|[,.\n()\[\]]|Tel\.?|Telefon|E-mail|RČ|OP|Datum|Kontakt|Nar\.?|Rodn[éě]|Číslo))',  # Lookahead
    re.UNICODE | re.IGNORECASE
)

# ADDRESS_REVERSE_RE - obrácený formát "Město, Ulice číslo" (pro texty jako "Praha 1, Washingtonova 1621/11")
# KRITICKÁ OPRAVA: Vyžaduje adresní prefix (jako ADDRESS_RE), aby se zabránilo false positive matchům
# Příklad false positive BEZ prefixu: "Dlužník potvrzuje, že uvedenou částku převzal v hotovosti dne 31"
#   → tento text by byl chybně detekován jako "město: Dlužník potvrzuje, ulice: že...dne, číslo: 31"
ADDRESS_REVERSE_RE = re.compile(
    r'(?<!\[)'                                       # Ne po '['
    r'(?:'                                           # Začátek prefixů (POVINNÉ!)
    r'(?:(?:trvale\s+)?bytem\s*:?\s*)|'             # "bytem" nebo "Bytem:"
    r'(?:(?:trvalé\s+)?bydlišt[eě]\s*:\s*)|'        # "trvalé bydliště:"
    r'(?:(?:sídlo(?:\s+podnikání)?|se\s+sídlem)\s*:\s*)|'  # "sídlo:" / "se sídlem:"
    r'(?:místo\s+(?:podnikání|výkonu\s+práce)\s*:?\s*)|'  # "Místo podnikání:" nebo "Místo výkonu práce" (volitelná :)
    r'(?:(?:adresa|trvalý\s+pobyt)\s*:\s*)|'       # "adresa:" / "trvalý pobyt:"
    r'(?:(?:v\s+ulic[ií]|na\s+adrese|v\s+dom[eě])\s+)'  # "v ulici " / "na adrese "
    r')'
    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'                         # Velké písmeno (začátek města)
    r'[a-záčďéěíňóřšťúůýž\s\d]{2,50}'               # Název města (Praha 1, České Budějovice) - GREEDY pro víceslovná města
    r',\s+'                                          # Čárka a mezera
    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'                         # Velké písmeno (začátek ulice)
    r'[a-záčďéěíňóřšťúůýž\s]{2,60}'                 # Název ulice - GREEDY pro víceslovné ulice
    r'\s+\d{1,4}(?:/\d{1,4})?'                      # Číslo domu (1621/11)
    r'(?=[\s,.]|$)',                                 # Zastaví se před mezerou, čárkou, tečkou nebo koncem
    re.UNICODE | re.IGNORECASE
)
ACCT_RE    = re.compile(r'\b(?:\d{1,6}-)?\d{2,10}/\d{4}\b')
BIRTHID_RE = re.compile(r'\b\d{6}\s*/\s*\d{3,4}\b')
IDCARD_RE  = re.compile(r'\b\d{6,9}/\d{3,4}\b|\b\d{9}\b|[A-Z]{2,3}[ \t]?\d{6,9}\b')
# KRITICKÁ OPRAVA: Rozšířený PHONE_RE pro detekci všech formátů včetně "420 777 111 222"
PHONE_RE   = re.compile(r'(?<!\d)(?:\+420|420|00420)?\s?\d{3}\s?\d{3}\s?\d{3}(?!\s*/\d{4})\b')
EMAIL_RE   = re.compile(r'[A-Za-z0-9._%+\-\u00C0-\u017F]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', re.UNICODE)
DATE_RE    = re.compile(r'\b\d{1,2}\.\s*\d{1,2}\.\s*\d{4}\b')

# DATE_WORDS_RE - detekuje datumy psané s českými názvy měsíců
# Příklady: "13. srpna 2025", "31. července 2025", "1. ledna 2024"
DATE_WORDS_RE = re.compile(
    r'\b(\d{1,2})\.\s+(ledna|února|března|dubna|května|června|července|srpna|září|října|listopadu|prosince)\s+(\d{4})\b',
    re.IGNORECASE | re.UNICODE
)

# LICENSE_PLATE_RE - detekuje české poznávací značky (SPZ/RZ)
# Formáty: "7AB 4567" (číslice + 2 písmena + mezera + 4 číslice)
#          "3M1 2345" (číslice + písmeno + číslice + mezera + 4 číslice)
#          "5AC 9845", "4BD 7654" atd.
LICENSE_PLATE_RE = re.compile(r'\b\d[A-Z]{1,2}\d?\s\d{4}\b')

# VIN_RE - detekuje VIN (Vehicle Identification Number)
# Formát: 17 znaků (velká písmena A-Z kromě I, O, Q + číslice 0-9)
# Příklad: TMBJK61Z3G0123456
VIN_RE = re.compile(r'\b[A-HJ-NPR-Z0-9]{17}\b')

# BIRTHPLACE_RE - detekuje místo narození pro GDPR compliance
# Příklad: "Místo narození: Brno", "Narozena v Praze"
BIRTHPLACE_RE = re.compile(
    r'(?:Místo\s+narození|Narozen[aáý]?\s+(?:v|ve)\s+|Rodiště)\s*:\s*'
    r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s\d]{2,50}?)'
    r'(?=\s*(?:$|[,.\n]|Rodn[éě]|RČ|OP|Občansk|Tel\.|Telefon|E-mail|Kontakt|Číslo|Datum|IČO|DIČ|Bydlišt|Bytem|Adresa))',
    re.IGNORECASE | re.UNICODE
)

STATUTE_RE = re.compile(r'\b(Sb\.?|zákon(a|u)?|zákon\s*č\.)\b', re.IGNORECASE)
PAIR_RE    = re.compile(r'(?<!\w)([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]{1,})\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]{1,})(?!\w)')
TITLES_RE  = re.compile(r'\b(Mgr|Ing|Dr|Ph\.?D|RNDr|MUDr|JUDr|PhDr|PaedDr|ThDr|RCDr|MVDr|DiS|Bc|BcA|MBA|LL\.?M|prof|doc|pan|paní|pán|slečna)\.?\s+', re.IGNORECASE)

# IČO a DIČ
ICO_RE     = re.compile(r'\bIČO\s*:?\s*(\d{8})\b', re.IGNORECASE)
DIC_RE     = re.compile(r'\bDIČ\s*:?\s*(CZ\d{8,10})\b', re.IGNORECASE)

# IBAN a BIC/SWIFT (GDPR - mezinárodní bankovní údaje)
# KRITICKÁ OPRAVA: IBAN detekce podporuje mezery (formát CZ65 0800 0000 0028 4756 3921)
IBAN_RE    = re.compile(r'\b(?:IBAN\s*:?\s*)?([A-Z]{2}\d{2}[A-Z0-9\s]{11,32})\b', re.IGNORECASE)
BIC_RE     = re.compile(r'\b([A-Z]{4}[A-Z]{2}[A-Z0-9]{2}(?:[A-Z0-9]{3})?)\b')  # BIC/SWIFT: 8 nebo 11 znaků

# KRITICKÁ OPRAVA: Platební karty (Visa, MasterCard, atd.)
# Podporuje 13-19 číslic s mezerami nebo pomlčkami
CARD_RE    = re.compile(
    r'\b(?:Platební\s+karta|Číslo\s+karty|Karta)\s*:?\s*(\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4,7})\b',
    re.IGNORECASE
)

# KRITICKÁ OPRAVA: IP adresy (IPv4)
IP_RE      = re.compile(r'\b(?:IP\s+adresa|IP)\s*:?\s*(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\b',
                       re.IGNORECASE)

# KRITICKÁ OPRAVA: Hesla a credentials
PASSWORD_RE = re.compile(r'\b(?:Initial\s+password|Password|Heslo)\s*:?\s*(\S+)', re.IGNORECASE)

# KRITICKÁ OPRAVA: Usernames, loginy, účty
USERNAME_RE = re.compile(r'\b(?:Login|Username|Uživatel|User|Account\s+ID)\s*:?\s*([A-Za-z0-9._\-@]+)', re.IGNORECASE)

# KRITICKÁ OPRAVA: API klíče a tajemství
API_KEY_RE  = re.compile(
    r'\b(?:AWS\s+Access\s+Key|AWS\s+Secret|API\s+Key|Stripe\s+API|SendGrid\s+API|Secret\s+Key)\s*:?\s*([A-Za-z0-9+/=]{20,})',
    re.IGNORECASE
)

# KRITICKÁ OPRAVA: Čísla pojištěnců (zdravotní identifikátory)
INSURANCE_ID_RE = re.compile(r'\b(?:Číslo\s+pojištěnce|Pojištěnec|Zdravotní\s+pojištění)\s*:?\s*(\d{9,10})', re.IGNORECASE)

# KRITICKÁ OPRAVA: RFID karty a badge
RFID_RE    = re.compile(r'\b(?:RFID\s+karta|RFID|Badge|ID\s+karta)\s*:?\s*([A-Za-z0-9\-_/]+)', re.IGNORECASE)

# KRITICKÁ OPRAVA: Řidičské průkazy
DRIVER_LICENSE_RE = re.compile(r'\b(?:Řidičský\s+průkaz|Řidičák)\s*.*?(?:č\.|číslo)\s*([A-Z0-9\s\-]+)', re.IGNORECASE)

# KRITICKÁ OPRAVA: Částky (aby se nezaměňovaly s telefony)
# Detekuje částky ve formátu "150 000 000", "1 500 000 Kč", atd.
AMOUNT_RE  = re.compile(r'\b(\d{1,3}(?:\s\d{3}){2,})\s*(?:Kč|EUR|USD|CZK)?\b')

# Osobní číslo zaměstnance
EMP_ID_RE  = re.compile(r'\b(?:osobn[íi]\s+č[íi]slo(?:\s+zaměstnance)?|zaměstnaneck[éeě]\s+č[íi]slo)\s*:?\s*(\d+)\b', re.IGNORECASE)

# Role-based jména (Jednatel: David Müller, Zaměstnanec: Nguyễn Thị Lan)
# Zachytává: "Role: Jméno Příjmení" nebo "Role: Jméno1 Jméno2 Příjmení" nebo "Role: Jméno "Přezdívka" Příjmení"
# Rozšířený Unicode rozsah pro vietnamská a jiná jména: \u00C0-\u024F (Latin Extended) + \u1E00-\u1EFF (Latin Extended Additional)
# Zastaví se před klíčovými slovy jako "Bytem:", "Bydliště:", "IČO:", atd.
ROLE_NAME_RE = re.compile(
    r'\b(Prodávající|Kupující|Zaměstnavatel|Zaměstnanec|Zaměstnavatelka|Zaměstnankyně|'
    r'Zhotovitel|Objednatel|Jednatel|Jednatelka|Makléř|Sv[eě]dek|'
    r'Pronaj[íi]matel|N[aá]jemce|Dlužn[íi]k|V[eě]řitel|Ručitel|Spoludlužn[íi]k|'
    r'Statut[aá]rn[íi]\s+z[aá]stupce)\s*:\s*'
    r'(?:(?:Mgr|Ing|Dr|Ph\.?D|RNDr|MUDr|JUDr|PhDr|PaedDr|ThDr|RCDr|MVDr|DiS|Bc|BcA|MBA|LL\.?M|prof|doc|pan|paní|pán|slečna)\.?\s+)?'  # volitelné tituly
    r'((?:[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽa-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]+\s+){0,2}[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽa-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]+)'  # Celé jméno (1-3 slova)
    r'(?=\s+(?:Bytem|Bydlišt[eě]|Sídlo|IČO|DIČ|Rodn[éě]|RČ|Nar\.|Tel\.|Telefon|Kontakt|E-mail|e-mail|OP|Občansk|Číslo|Datum|$))',  # Zastaví se před klíčovými slovy
    re.IGNORECASE | re.UNICODE
)

# Jména s přezdívkami (Martin "Marty" Král)
NICKNAME_RE = re.compile(
    r'(?<!\w)([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,20})\s+"([^"]{1,20})"\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,20})(?!\w)',
    re.UNICODE
)

# Samostatné přezdívky v textu (dále jen "Marty", "Marty", apod.)
STANDALONE_NICKNAME_RE = re.compile(
    r'\(dále\s+jen\s+"([^"]{1,20})"\)',
    re.UNICODE | re.IGNORECASE
)

# Multi-token foreign names (Nguyễn Thị Lan - dvě křestní jména + příjmení)
MULTI_TOKEN_NAME_RE = re.compile(
    r'(?<!\w)([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,15})\s+'
    r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,15})\s+'
    r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,15})(?!\w)',
    re.UNICODE
)

CTX_OP     = re.compile(r'\b(OP|Číslo\s+OP|číslo\s+OP|občansk(ý|ého|ému|ém|ým)|průkaz|č\.\s*OP)\b', re.IGNORECASE)
CTX_BIRTH  = re.compile(r'\b(rodn[ée]\s*č[íi]slo|r\.?\s*č\.?|RČ|rodn[ée])\b', re.IGNORECASE)
CTX_BANK   = re.compile(r'\b(účet|účtu|účtem|Bankovní\s+účet|bankovní\s+účet|veden[eya].*u|banka|banky|IBAN|číslo\s+účtu)\b', re.IGNORECASE)
CTX_PERSON = re.compile(
    r'(nar\.|narozen|rodn[ée]\s*č[íi]slo|RČ|bytem|trval[é]\s*bydlišt[ěi]|'
    r'(e-?mail)|tel\.?|telefon|č\.\s*účtu|IBAN|SPZ|Mgr\.|Ing\.|Bc\.|PhDr\.|JUDr\.)',
    re.IGNORECASE
)
CTX_ROLE   = re.compile(r'\b(pronaj[ií]matel|n[aá]jemce|dlu[zž]n[ií]k|v[eě]řitel|objednatel|zhotovitel|zam[eě]stnanec|zam[eě]stnavatel|ručitel|spoludlu[zž]n[ií]k|jednatel|statut[aá]rn[ií]\s+z[aá]stupce|sv[eě]dek)\b', re.IGNORECASE)
CTX_LABEL  = re.compile(r'j[mn][eě]no\s*(,|a)?\s*př[ií]jmen[ií]', re.IGNORECASE)

def looks_like_firstname(token: str) -> bool:
    if not token or not token[0].isupper(): return False
    norm = normalize_for_matching(token)
    if norm in CZECH_FIRST_NAMES: return True
    return any([
        norm.endswith('ek'), norm.endswith('el'), norm.endswith('os'),
        norm.endswith('as'), norm.endswith('an'), norm.endswith('en'),
        norm.endswith('a') and len(norm) > 3,
    ])

# =============== Anonymizer ===============
class Anonymizer:
    def __init__(self, verbose=False):
        self.verbose = verbose
        self.counter = defaultdict(int)
        self.tag_map = defaultdict(list)
        self.value_to_tag = {}
        self.person_index = {}
        self.canonical_persons = []
        self.person_variants = {}
        self.source_text = ""

    def _get_or_create_tag(self, cat: str, value: str) -> str:
        norm_val = ' '.join(value.split())
        lookup_key = f"{cat}:{norm_val}"
        if lookup_key in self.value_to_tag:
            return self.value_to_tag[lookup_key]
        self.counter[cat] += 1
        tag = f'[[{cat}_{self.counter[cat]}]]'
        self.value_to_tag[lookup_key] = tag
        self._record_value(tag, value)
        return tag

    def _record_value(self, tag: str, value: str):
        # Normalize: odstranění leading/trailing mezer a vícenásobných mezer
        value = re.sub(r'\s+', ' ', value).strip()
        if not value:
            return

        # Pro DATE tagy ukládat vždy (normalizované hodnoty nemusí být v původním textu)
        if tag.startswith('[[DATE_'):
            if value not in self.tag_map[tag]:
                self.tag_map[tag].append(value)
        # Pro ostatní tagy kontrolovat, zda hodnota existuje v původním textu
        elif re.search(r'(?<!\w)'+re.escape(value)+r'(?!\w)', self.source_text):
            if value not in self.tag_map[tag]:
                self.tag_map[tag].append(value)

    def _ensure_person_tag(self, first_nom: str, last_nom: str) -> str:
        key = (normalize_for_matching(first_nom), normalize_for_matching(last_nom))
        if key in self.person_index:
            return self.person_index[key]
        tag = self._get_or_create_tag('PERSON', f'{first_nom} {last_nom}')
        self.person_index[key] = tag
        self.canonical_persons.append({'first': first_nom, 'last': last_nom, 'tag': tag})

        # KRITICKÁ OPRAVA: Zajisti, že kanonická forma (nominativ) je VŽDY první v tag_map
        # i když není přímo v původním textu (může být jen pádová forma)
        canonical_full = f'{first_nom} {last_nom}'
        if canonical_full not in self.tag_map[tag]:
            # Vlož kanonickou formu na PRVNÍ místo
            self.tag_map[tag].insert(0, canonical_full)

        fvars = variants_for_first(first_nom)
        svars = variants_for_surname(last_nom)
        self.person_variants[tag] = {f'{f} {s}' for f in fvars for s in svars}
        return tag

    def _extract_persons_to_index(self, text: str):
        # FÁZE 0a: Konservativní detekce jmen po specifických rolích (Jednatel:, Zaměstnanec:, atd.)
        # Podporuje 2-3 slovná jména (David Müller, Nguyễn Thị Lan)
        simple_role_re = re.compile(
            r'\b(Jednatel|Jednatelka|Zaměstnanec|Zaměstnankyně|Dlužn[íi]k|V[eě]řitel|Prodávající|Kupující)\s*:\s*'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,20})'  # První jméno
            r'(?:\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,20}))?'  # Volitelné prostřední jméno
            r'\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\u00C0-\u024F\u1E00-\u1EFF][a-záčďéěíňóřšťúůýž\u00C0-\u024F\u1E00-\u1EFF]{1,20})'  # Příjmení (poslední slovo)
            r'(?=\s+(?:Bytem|Bydlišt|Sídlo|E-mail|Tel|Kontakt|$))',  # Zastaví se před klíčovými slovy
            re.IGNORECASE | re.UNICODE
        )

        for m in simple_role_re.finditer(text):
            first_part = m.group(2)
            middle_part = m.group(3)  # může být None
            surname = m.group(4)

            # Pokud je prostřední jméno, zkombinuj ho s první částí
            if middle_part:
                f_nom = f"{first_part} {middle_part}"
            else:
                f_nom = first_part

            # Kontrola blacklistu - ale dovolíme "nový/nová" jako příjmení pokud jsou po roli
            # (Adam Nový, Petra Nová jsou běžná jména i když "nový" je adjektivum)
            surname_norm = normalize_for_matching(surname)
            fname_norm = normalize_for_matching(f_nom)

            # Skip common blacklisted words, but allow "novy/nova" as it's also a surname
            if surname_norm in SURNAME_BLACKLIST and surname_norm not in ('novy', 'nova', 'nove'):
                continue
            if fname_norm in SURNAME_BLACKLIST and fname_norm not in ('novy', 'nova', 'nove'):
                continue

            f_nom_inferred = infer_first_name_nominative(first_part, surname) or f_nom
            l_nom = infer_surname_nominative(surname)

            self._ensure_person_tag(f_nom, l_nom)

        # FÁZE 0b: Detekce jmen s přezdívkami (Martin "Marty" Král)
        for m in NICKNAME_RE.finditer(text):
            first_name = m.group(1)
            nickname = m.group(2)
            surname = m.group(3)

            # Kontrola blacklistu
            if normalize_for_matching(surname) in SURNAME_BLACKLIST:
                continue
            if normalize_for_matching(first_name) in SURNAME_BLACKLIST:
                continue

            f_nom = infer_first_name_nominative(first_name, surname) or first_name
            l_nom = infer_surname_nominative(surname)

            self._ensure_person_tag(f_nom, l_nom)

        # FÁZE 1: Standardní dvojice (Křestní Příjmení)
        text_no_titles = TITLES_RE.sub('', text)
        for m in PAIR_RE.finditer(text_no_titles):
            s, e = m.span()
            f_tok, l_tok = m.group(1), m.group(2)

            if f_tok.lower() in ROLE_STOP or l_tok.lower() in ROLE_STOP:
                continue
            if normalize_for_matching(l_tok) in SURNAME_BLACKLIST:
                continue
            if normalize_for_matching(f_tok) in SURNAME_BLACKLIST:
                continue
            
            pre = text[max(0, s-80):s]
            post = text[e:e+80]

            # KRITICKÁ OPRAVA: Organizace a firmy
            # Pokud je za jménem "a.s.", "s.r.o.", "spol.", atd., je to firma, ne osoba
            if re.search(r'\s+(a\.s\.|s\.r\.o\.|spol\.|v\.o\.s\.|o\.p\.s\.|o\.s\.|z\.s\.)', post, re.IGNORECASE):
                continue

            # Pokud je před jménem "Oddělení:", "Instituce:", "Společnost:", je to organizace
            if re.search(r'\b(Oddělení|Instituce|Společnost|Korporace|Organizace|Firma)\s*:\s*$', pre, re.IGNORECASE):
                continue

            if re.search(r'\b(výrobce|model|značka|inventář|výrobek|položk)', pre+post, re.IGNORECASE):
                if (normalize_for_matching(f_tok) in SURNAME_BLACKLIST or
                    normalize_for_matching(l_tok) in SURNAME_BLACKLIST):
                    continue

            f_nom = infer_first_name_nominative(f_tok, l_tok) or f_tok
            l_nom = infer_surname_nominative(l_tok)

            if normalize_for_matching(f_nom) in CZECH_FIRST_NAMES:
                self._ensure_person_tag(f_nom, l_nom)
                continue

            pre = text[max(0, s-160):s]
            post = text[e:e+160]
            has_ctx = CTX_PERSON.search(pre+post) or CTX_ROLE.search(pre+post) or CTX_LABEL.search(pre+post)
            if (has_ctx
                and f_tok[:1].isupper() and l_tok[:1].isupper()
                and looks_like_firstname(f_tok)
                and f_tok.lower() not in ROLE_STOP and l_tok.lower() not in ROLE_STOP):
                self._ensure_person_tag(f_nom, l_nom)

    def _apply_known_people(self, text: str) -> str:
        # FÁZE 0b: Nahrazení jmen s přezdívkami (Martin "Marty" Král)
        def nickname_repl(m):
            first_name = m.group(1)
            nickname = m.group(2)
            surname = m.group(3)

            # Kontrola blacklistu
            if normalize_for_matching(surname) in SURNAME_BLACKLIST:
                return m.group(0)
            if normalize_for_matching(first_name) in SURNAME_BLACKLIST:
                return m.group(0)

            f_nom = infer_first_name_nominative(first_name, surname) or first_name
            l_nom = infer_surname_nominative(surname)

            key = (normalize_for_matching(f_nom), normalize_for_matching(l_nom))
            if key in self.person_index:
                tag = self.person_index[key]
                full_match = m.group(0)
                self._record_value(tag, full_match)
                return preserve_case(full_match, tag)

            return m.group(0)

        text = NICKNAME_RE.sub(nickname_repl, text)

        # FÁZE 1: Nahrazení plných jmen (křestní + příjmení)
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])
            for pat in sorted(self.person_variants[tag], key=len, reverse=True):
                rx = re.compile(r'(?<!\w)'+re.escape(pat)+r'(?!\w)', re.IGNORECASE)
                def repl(m):
                    surf = m.group(0)
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)
                text = rx.sub(repl, text)

            # FÁZE 2: Nahrazení přivlastňovacích přídavných jmen (Novákův, Janin)
            first_low, last_low = p['first'].lower(), p['last'].lower()
            poss = set()
            if first_low.endswith('a'):
                stem = p['first'][:-1]
                poss |= {stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných']}
                if stem.endswith('tr'):
                    poss |= {stem[:-1]+'ř'+s for s in ['in','ina','iny','iné','inu','inou','iným','iných']}
            else:
                poss |= {p['first']+'ův'} | {p['first']+'ov'+s for s in ['a','o','y','ě','ým','ých']}
            if not last_low.endswith('ová'):
                poss |= {p['last']+'ův'} | {p['last']+'ov'+s for s in ['a','o','y','ě','ým','ých']}
            for token in sorted(list(poss), key=len, reverse=True):
                rx = re.compile(r'(?<!\w)'+re.escape(token)+r'(?!\w)', re.IGNORECASE)
                def repl2(m):
                    surf = m.group(0)
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)
                text = rx.sub(repl2, text)

        # FÁZE 3: Nahrazení samostatných příjmení (bez křestního jména)
        # Příklad: "Horváthová pronajímá Procházkovi byt. Procházka platí Horváthové nájemné."
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])

            # Generuj všechny pádové varianty příjmení
            surname_variants = variants_for_surname(p['last'])

            # Také přidej varianty křestního jména pro kontrolu
            first_variants = variants_for_first(p['first'])
            # Normalizuj křestní jména pro kontrolu (lowercase pro case-insensitive matching)
            first_variants_lower = {fv.lower() for fv in first_variants if fv}

            for surname_var in sorted(surname_variants, key=len, reverse=True):
                if not surname_var or len(surname_var) < 2:
                    continue

                # Jednoduchý regex pro nalezení příjmení jako samostatného slova
                rx = re.compile(r'(?<!\w)' + re.escape(surname_var) + r'(?!\w)', re.IGNORECASE)

                # Použijeme callback funkci, která zkontroluje kontext
                def repl3_with_context(m):
                    surf = m.group(0)
                    start_pos = m.start()
                    end_pos = m.end()

                    # DŮLEŽITÉ: Přeskoč příjmení uvnitř "(rozená Xxx)" nebo "(dříve Xxx)"
                    # Toto zabraňuje kolizi tagů (např. "(rozená Nová)" nesloučí s "Adam Nový")
                    context_before_wide = text[max(0, start_pos-30):start_pos]
                    if re.search(r'\((?:rozená|roz\.?|dříve)\s+(?:[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]\w+\s+)?$', context_before_wide, re.IGNORECASE):
                        # Příjmení je uvnitř "(rozená ...)" - přeskoč!
                        return surf

                    # Zkontroluj 50 znaků před a 50 znaků po
                    context_before = text[max(0, start_pos-50):start_pos]
                    context_after = text[end_pos:min(len(text), end_pos+50)]

                    # Extrahuj poslední slovo před a první slovo po
                    words_before = re.findall(r'\b\w+\b', context_before)
                    words_after = re.findall(r'\b\w+\b', context_after)

                    # Pokud poslední slovo je oslovení/titul (Paní, Pan, MUDr., atd.), IGNORUJ ho
                    titles_and_salutations = {'pan', 'paní', 'pani', 'pana', 'panu', 'mudr', 'ing', 'mgr', 'judr', 'bc', 'doc', 'prof'}
                    if words_before and words_before[-1].lower() in titles_and_salutations:
                        # Odstraň titul ze seznamu slov před
                        words_before = words_before[:-1]

                    # Pokud poslední slovo před příjmením je křestní jméno, NENAHRAZUJ
                    if words_before and words_before[-1].lower() in first_variants_lower:
                        return surf  # Nech to být (je to součást celého jména)

                    # Pokud první slovo po příjmení je křestní jméno, NENAHRAZUJ
                    if words_after and words_after[0].lower() in first_variants_lower:
                        return surf  # Nech to být

                    # Jinak je to samostatné příjmení → anonymizuj
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)

                text = rx.sub(repl3_with_context, text)

        # FÁZE 3b: Nahrazení slov z křestního jména (pro vietnamská/asijská jména kde je příjmení první)
        # Například: "Paní Nguyễn" kde "Nguyễn" je technicky v 'first', ale je to příjmení
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])

            # Rozděl křestní jméno na slova (např. "Nguyễn Thị" -> ["Nguyễn", "Thị"])
            first_words = p['first'].split()

            # Pro každé slovo z křestního jména (kromě velmi krátkých)
            for word in first_words:
                if len(word) < 3:  # Přeskoč velmi krátká slova
                    continue

                # Pokud slovo vypadá jako příjmení (velké písmeno na začátku, delší než 3 znaky)
                if word[0].isupper() and len(word) >= 3:
                    rx = re.compile(r'(?<!\w)' + re.escape(word) + r'(?!\w)', re.IGNORECASE)

                    def repl3b(m):
                        surf = m.group(0)
                        start_pos = m.start()

                        # Zkontroluj kontext
                        context_before = text[max(0, start_pos-50):start_pos]
                        words_before = re.findall(r'\b\w+\b', context_before)

                        # Pokud je před slovem "Paní/Pan" nebo jiný titul, anonymizuj
                        titles = {'pan', 'paní', 'pani', 'pana', 'panu', 'panem', 'mudr', 'ing', 'mgr'}
                        if words_before and words_before[-1].lower() in titles:
                            self._record_value(tag, surf)
                            return preserve_case(surf, tag)

                        # Jinak nech to být
                        return surf

                    text = rx.sub(repl3b, text)

        # FÁZE 3.5: Speciální handler pro "(rozená Xxx)" / "(roz. Xxx)" / "(dříve Xxx)"
        # DŮLEŽITÉ: Musí být PO FÁZÍ 3 (aby už byly samostatná příjmení nahrazená jako [[PERSON_*]])
        # aby handler mohl najít předchozí [[PERSON_*]] tag ve větě
        MAIDEN_NAME_RE = re.compile(
            r'\((?:rozená|roz\.?|dříve)\s+(?:([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\s+)?([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\)',
            re.IGNORECASE | re.UNICODE
        )

        def maiden_name_repl(m):
            full_match = m.group(0)
            s, e = m.span()
            first_name = m.group(1)  # Může být None
            maiden_surname = m.group(2)
            keyword = m.group(0).split()[0][1:]  # Extrahuj "rozená" nebo "dříve" z "(rozená"

            # Hledej předchozí PERSON tag ve větě (do 200 znaků zpět)
            pre = text[max(0, s-200):s]
            # Hledej nejbližší PERSON tag (pozpátku = od konce = nejbližší)
            person_tags = list(re.finditer(r'\[\[PERSON_\d+\]\]', pre))

            if person_tags:
                person_tag = person_tags[-1].group(0)  # Poslední = nejbližší
                # Přidej rodné jméno k hodnotám tohoto tagu
                self._record_value(person_tag, full_match)

                # Vytvoř anonymizovanou verzi se zachovaným klíčovým slovem
                return f'({keyword} {person_tag})'  # "(rozená [[PERSON_X]])" nebo "(dříve [[PERSON_X]])"

            # Pokud nenajdeme předchozí PERSON tag, nech to být
            return full_match

        text = MAIDEN_NAME_RE.sub(maiden_name_repl, text)

        # FÁZE 3.7: Nahrazení samostatných křestních jmen (bez příjmení)
        # Příklad: "Petra uhradí Martinovi částku" → "[[PERSON_16]] uhradí [[PERSON_5]] částku"
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])

            # Generuj všechny pádové varianty křestního jména
            first_variants = variants_for_first(p['first'])

            # Také přidej varianty příjmení pro kontrolu
            surname_variants = variants_for_surname(p['last'])
            surname_variants_lower = {sv.lower() for sv in surname_variants if sv}

            for first_var in sorted(first_variants, key=len, reverse=True):
                if not first_var or len(first_var) < 2:
                    continue

                # Regex pro nalezení křestního jména jako samostatného slova
                rx = re.compile(r'(?<!\w)' + re.escape(first_var) + r'(?!\w)', re.IGNORECASE)

                def repl_first_with_context(m):
                    surf = m.group(0)
                    start_pos = m.start()
                    end_pos = m.end()

                    # Zkontroluj kontext (50 znaků před a po)
                    context_before = text[max(0, start_pos-50):start_pos]
                    context_after = text[end_pos:min(len(text), end_pos+50)]

                    # Extrahuj slova kolem
                    words_before = re.findall(r'\b\w+\b', context_before)
                    words_after = re.findall(r'\b\w+\b', context_after)

                    # Pokud následuje nebo předchází příjmení této osoby, NENAHRAZUJ
                    # (je to součást plného jména, bude nahrazeno v FÁZI 1)
                    if words_after and words_after[0].lower() in surname_variants_lower:
                        return surf  # Plné jméno
                    if words_before and words_before[-1].lower() in surname_variants_lower:
                        return surf  # Plné jméno

                    # DŮLEŽITÉ: Pokud existuje v širším kontextu (200 znaků zpět) PERSON tag
                    # který obsahuje toto křestní jméno, použij TEN tag místo tohoto!
                    # Toto řeší problém disambiguation (Petra = Petr Novotný vs. Petra Beránková)
                    wide_context = text[max(0, start_pos-200):start_pos]
                    nearby_person_tags = list(re.finditer(r'\[\[PERSON_\d+\]\]', wide_context))

                    if nearby_person_tags:
                        # Najdi poslední (= nejbližší) PERSON tag
                        nearest_tag = nearby_person_tags[-1].group(0)

                        # Zkontroluj, jestli tento tag obsahuje variantu našeho křestního jména
                        if nearest_tag in self.tag_map:
                            for val in self.tag_map[nearest_tag]:
                                # Extrahuj křestní jméno z hodnoty (první slovo)
                                val_words = val.split()
                                if val_words:
                                    val_first = val_words[0]
                                    # Pokud první slovo matchuje náš surf (case-insensitive)
                                    if val_first.lower() == surf.lower():
                                        # Použij nejbližší tag!
                                        self._record_value(nearest_tag, surf)
                                        return preserve_case(surf, nearest_tag)

                    # Jinak je to samostatné křestní jméno → anonymizuj s tímto tageem
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)

                text = rx.sub(repl_first_with_context, text)

        # FÁZE 4: Nahrazení samostatných přezdívek v textu (dále jen "Marty")
        # Propojíme je se známými osobami na základě přezdívky
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])

            # Zkontroluj, zda osoba má přezdívku v hodnotách
            nicknames = set()
            for val in self.tag_map.get(tag, []):
                # Hledej přezdívky ve formátu 'Name "Nickname" Surname'
                nick_match = NICKNAME_RE.search(val)
                if nick_match:
                    nicknames.add(nick_match.group(2).lower())

            # Anonymizuj standalone přezdívky
            for nickname in nicknames:
                # Pattern: (dále jen "nickname")
                pattern = re.compile(r'\(dále\s+jen\s+"' + re.escape(nickname) + r'"\)', re.IGNORECASE)
                def nickname_standalone_repl(m):
                    self._record_value(tag, m.group(0))
                    return f'(dále jen "{tag}")'
                text = pattern.sub(nickname_standalone_repl, text)

        return text

    def _replace_remaining_people(self, text: str) -> str:
        text_no_titles = TITLES_RE.sub('', text)
        offset = 0
        for m in list(PAIR_RE.finditer(text_no_titles)):
            s, e = m.start()+offset, m.end()+offset
            seg = text[s:e]
            if seg.startswith('[[') and seg.endswith(']]'):
                continue
            f_tok, l_tok = m.group(1), m.group(2)

            if f_tok.lower() in ROLE_STOP or l_tok.lower() in ROLE_STOP:
                continue
            if normalize_for_matching(l_tok) in SURNAME_BLACKLIST:
                continue
            # KRITICKÁ OPRAVA: Kontrola křestního jména proti blacklistu
            # Zabránit detekci "Položka Stav" jako jméno
            if normalize_for_matching(f_tok) in SURNAME_BLACKLIST:
                continue

            f_nom = infer_first_name_nominative(f_tok, l_tok) or f_tok
            pre = text[max(0, s-160):s]
            post = text[e:e+160]
            has_ctx = CTX_PERSON.search(pre+post) or CTX_ROLE.search(pre+post) or CTX_LABEL.search(pre+post)

            if (normalize_for_matching(f_nom) not in CZECH_FIRST_NAMES
                and not (has_ctx and looks_like_firstname(f_tok))):
                continue

            l_nom = infer_surname_nominative(l_tok)
            tag = self._ensure_person_tag(f_nom, l_nom)
            before = text
            text = text[:s] + preserve_case(seg, tag) + text[e:]
            self._record_value(tag, seg)
            offset += len(text) - len(before)
        return text

    def _is_statute(self, text: str, s: int, e: int) -> bool:
        pre = text[max(0, s-20):s]
        post = text[e:e+10]
        return bool(STATUTE_RE.search(pre) or STATUTE_RE.search(post))

    def _replace_entity(self, text: str, rx: re.Pattern, cat: str) -> str:
        def repl(m):
            v = m.group(0)
            tag = self._get_or_create_tag(cat, v)
            self._record_value(tag, v)
            return tag
        return rx.sub(repl, text)

    def anonymize_entities(self, text: str) -> str:
        # KRITICKÁ OPRAVA: E-MAILY MUSÍ BÝT ÚPLNĚ PRVNÍ!
        # Jinak se jména v e-mailech (např. "martina.horáková@example.com") nahradí jako osoby
        # a zbyde "[[PERSON]].horáková@example.com"
        text = self._replace_entity(text, EMAIL_RE, 'EMAIL')

        # SPECIÁLNÍ PŘÍPAD: "Jméno Příjmení, bytem Adresa" (např. v Svědcích)
        # Musí být PŘED zpracováním adres a osob!
        PERSON_BYTEM_ADDRESS_RE = re.compile(
            r'(?<!\[)'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+(?:\s+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)*)'  # Jméno (+ příjmení)
            r',\s+'
            r'(bytem\s+)'  # "bytem " (zachovat)
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s]+\s+\d{1,4}(?:/\d{1,4})?)',  # Adresa bez města
            re.IGNORECASE | re.UNICODE
        )

        def person_bytem_repl(m):
            person_name = m.group(1).strip()
            bytem_prefix = m.group(2)
            address = m.group(3).strip()

            # Rozděl jméno na křestní jméno a příjmení
            name_parts = person_name.split()
            if len(name_parts) >= 2:
                first_name = name_parts[0]
                last_name = ' '.join(name_parts[1:])
            else:
                first_name = person_name
                last_name = person_name

            # Vytvoř tagy
            person_tag = self._ensure_person_tag(first_name, last_name)
            self._record_value(person_tag, person_name)

            address_tag = self._get_or_create_tag('ADDRESS', address)
            self._record_value(address_tag, address)

            return f'{person_tag}, {bytem_prefix}{address_tag}'

        text = PERSON_BYTEM_ADDRESS_RE.sub(person_bytem_repl, text)

        # DŮLEŽITÉ: Adresy DRUHÉ! (po e-mailech, ale před osobami)
        # Jinak "Novákova 45" se detekuje jako jméno
        def addr_repl(m):
            full_match = m.group(0)
            v = full_match.strip()
            s, e = m.span()
            pre = text[max(0, s-20):s]

            # DŮLEŽITÉ: Pokud je před matchem "OP:", je to občanský průkaz, ne adresa!
            # Např: "OP: AB 456789, vydán 12" by se jinak detekoval jako adresa
            if re.search(r'\bOP\s*:\s*$', pre, re.IGNORECASE):
                return full_match  # Neanonymizuj, nechej pro IDCARD_RE

            # DŮLEŽITÉ: Pokud match obsahuje ", bytem", může to být "Jméno Příjmení, bytem Adresa"
            # Např: "Martin Novák, bytem Nová Ves 78" by se jinak detekoval jako adresa
            if re.search(r',\s+bytem\s+', v, re.IGNORECASE):
                return full_match  # Neanonymizuj, nechej pro separátní zpracování jména a adresy

            # Zachytit prefix PŘED odstraněním (pro zachování v textu)
            # Dvojtečka je volitelná pro případy jako "Článek II - Místo výkonu práce Praha 1..."
            prefix_match = re.match(r'^(Trvalé\s+bydliště|Bydliště|(?:Trvale\s+)?[Bb]ytem|Adresa|Místo\s+(?:podnikání|výkonu\s+práce)|Sídlo\s+podnikání|Se\s+sídlem|Sídlo|Trvalý\s+pobyt)\s*:?\s*', v, flags=re.IGNORECASE)
            prefix = prefix_match.group(0) if prefix_match else ''

            # Odstranění běžných prefixů adres (s dvojtečkou i bez)
            # DŮLEŽITÉ: "trvale bytem" musí být před samotným "bytem" (delší vzor má přednost)
            v = re.sub(r'^(Trvalé\s+bydliště|Bydliště|[Tt]rvale\s+bytem|[Bb]ytem|Adresa|Místo\s+(?:podnikání|výkonu\s+práce)|Sídlo\s+podnikání|Se\s+sídlem|Sídlo|Trvalý\s+pobyt)\s*:?\s*', '', v, flags=re.IGNORECASE)

            # Odstranění kontextových/narrativních frází (např. "NP domu na adrese", "v 2. NP domu", "domu na adrese")
            # Zachytí různé varianty: "NP domu na adrese", "v 1. NP domu", "v domě na adrese", "na adrese", "v ulici"
            v = re.sub(r'^(?:(?:v\s+)?(?:\d+\.)?\s*NP\s+)?(?:domu\s+)?(?:na\s+adrese|v\s+dom[eě]|v\s+ulic[ií])\s+', '', v, flags=re.IGNORECASE)

            # Odstranění závorek a všeho v nich
            v = re.sub(r'\s*\(.*?\)\s*', ' ', v, flags=re.IGNORECASE)
            v = re.sub(r'\s*\(dále\s+jen.*$', '', v, flags=re.IGNORECASE)

            # Odstranění přebytečných mezer
            v = re.sub(r'\s+', ' ', v)
            v = v.strip()

            if not v:
                return full_match
            tag = self._get_or_create_tag('ADDRESS', v)
            self._record_value(tag, v)

            # Vrátit prefix + tag (zachování kontextu)
            # Pokud prefix neobsahuje dvojtečku, přidej ji pro čitelnost
            if prefix and not prefix.rstrip().endswith(':'):
                prefix = prefix.rstrip() + ': '

            return prefix + tag

        # KRITICKÁ OPRAVA: Adresy s PSČ BEZ/S prefixem
        # Musí být PRVNÍ, protože je nejspecifičtější (vyžaduje PSČ)
        # Group 1: prefix (může být prázdný), Group 2: adresa
        def addr_with_zip_repl(m):
            prefix = m.group(1) if m.group(1) else ''  # Prefix (v 2. NP domu na adrese)
            v = m.group(2).strip()  # Adresa bez prefixu

            # Odstranění běžných prefixů z hodnoty (pro mapu)
            v_clean = re.sub(r'^(Trvalé\s+bydliště|Bydliště|[Tt]rvale\s+bytem|[Bb]ytem|Adresa|Místo\s+(?:podnikání|výkonu\s+práce)|Sídlo\s+podnikání|Se\s+sídlem|Sídlo|Trvalý\s+pobyt)\s*:?\s*', '', v, flags=re.IGNORECASE)

            # Odstranění závorek
            v_clean = re.sub(r'\s*\(.*?\)\s*', ' ', v_clean, flags=re.IGNORECASE)
            v_clean = re.sub(r'\s+', ' ', v_clean).strip()

            if not v_clean:
                return m.group(0)

            tag = self._get_or_create_tag('ADDRESS', v_clean)
            self._record_value(tag, v_clean)
            # DŮLEŽITÉ: Vracíme prefix + tag, aby se kontext zachoval
            return prefix + tag
        text = ADDRESS_WITH_ZIP_RE.sub(addr_with_zip_repl, text)

        # Pak standardní formát "Ulice číslo, Město" S PREFIXEM
        text = ADDRESS_RE.sub(addr_repl, text)

        # Pak obrácený formát "Město, Ulice číslo" (např. "Praha 1, Washingtonova 1621/11")
        text = ADDRESS_REVERSE_RE.sub(addr_repl, text)

        # GDPR: SPZ/RZ (poznávací značky) jsou osobní identifikátory vozidla
        text = self._replace_entity(text, LICENSE_PLATE_RE, 'LICENSE_PLATE')

        # GDPR: VIN (Vehicle Identification Number) - 17-znakový kód vozidla
        text = self._replace_entity(text, VIN_RE, 'VIN')

        # POZNÁMKA: E-maily jsou zpracovány na ZAČÁTKU funkce (před adresami a osobami)

        # Datumy - normalizovat na DD.MM.RRRR formát
        def date_repl(m):
            original = m.group(0)  # Původní hodnota z textu
            # Parse date: "10.4.2025" → "10.04.2025", "23.09.1985" → "23.09.1985"
            parts = re.split(r'[.\s]+', original.strip())
            if len(parts) == 3:
                day = parts[0].zfill(2)
                month = parts[1].zfill(2)
                year = parts[2]
                normalized = f'{day}.{month}.{year}'
            else:
                normalized = original  # Fallback

            tag = self._get_or_create_tag('DATE', normalized)
            self._record_value(tag, normalized)  # OPRAVA: Ukládat normalizovanou formu pro konzistenci
            return tag

        text = DATE_RE.sub(date_repl, text)

        # Datumy psané slovy ("13. srpna 2025") - konvertovat na DD.MM.RRRR
        MONTH_MAP = {
            'ledna': '01', 'února': '02', 'března': '03', 'dubna': '04',
            'května': '05', 'června': '06', 'července': '07', 'srpna': '08',
            'září': '09', 'října': '10', 'listopadu': '11', 'prosince': '12'
        }
        def date_words_repl(m):
            original = m.group(0)  # Původní hodnota ("13. srpna 2025")
            day = m.group(1).zfill(2)  # 1 → 01
            month_name = m.group(2).lower()
            year = m.group(3)

            month_num = MONTH_MAP.get(month_name, '??')
            normalized = f'{day}.{month_num}.{year}'

            tag = self._get_or_create_tag('DATE', normalized)
            self._record_value(tag, normalized)  # OPRAVA: Ukládat normalizovanou formu pro eliminaci duplicit
            return tag

        text = DATE_WORDS_RE.sub(date_words_repl, text)

        # GDPR: Místo narození (toponyma jsou PII)
        def birthplace_repl(m):
            full_match = m.group(0)
            place = m.group(1).strip()

            # Zachytit prefix PŘED místem (pro zachování v textu)
            prefix_match = re.match(r'^(.*?:\s*)', full_match, re.IGNORECASE)
            prefix = prefix_match.group(1) if prefix_match else ''

            # Vytvoř tag pro místo
            tag = self._get_or_create_tag('PLACE', place)
            self._record_value(tag, place)

            # Vrátit prefix + tag
            return prefix + tag

        text = BIRTHPLACE_RE.sub(birthplace_repl, text)

        # KRITICKÁ OPRAVA: Telefony PŘED částkami! (jinak "420 777 111 222" matchuje jako částka)
        def phone_repl(m):
            v = m.group(0)
            s, e = m.span()
            pre = text[max(0, s-15):s]
            if re.search(r'(OP|občansk\w+|č\.\s*OP)', pre, re.IGNORECASE):
                tag = self._get_or_create_tag('ID_CARD', v)
                self._record_value(tag, v)
                return tag
            if re.match(r'^\s*/\d{4}', text[e:e+6]):
                return v
            tag = self._get_or_create_tag('PHONE', v)
            self._record_value(tag, v)
            return tag
        text = PHONE_RE.sub(phone_repl, text)

        # Částky AŽ PO telefonech (aby čísla jako "420 777 111 222" byla správně telefony)
        def amount_repl(m):
            v = m.group(1)
            tag = self._get_or_create_tag('AMOUNT', v)
            self._record_value(tag, v)
            # Vrátit celý match (včetně měny pokud je)
            return m.group(0).replace(v, tag)
        text = AMOUNT_RE.sub(amount_repl, text)

        def acct_like(m):
            s, e = m.span()
            if self._is_statute(text, s, e):
                return m.group(0)
            raw = m.group(0)

            # KRITICKÁ POLITIKA: Shape má přednost před labelem!
            # Pokud má tvar RČ (6 číslic / 3-4 číslice) → neanonymizuj zde
            # Nech to pro BIRTHID_RE který ho správně označí jako BIRTH_ID
            if re.match(r'^\d{6}/\d{3,4}$', raw):
                return raw  # Vrátit bez změny, bude zpracováno jako BIRTH_ID

            pre = text[max(0, s-30):s]
            post = text[e:e+30]

            # DŮLEŽITÉ: Pokud je to RČ (rodné číslo), NEANONYMIZUJ zde
            # Nech to pro BIRTHID_RE který běží později
            if CTX_BIRTH.search(pre+post):
                return raw  # Vrátit bez změny, bude zpracováno jako BIRTH_ID

            parts = raw.split('/')
            if len(parts) == 2:
                main_part = parts[0].replace('-', '')
                bank_code = parts[1]

                if len(main_part) >= 7 and len(bank_code) == 4:
                    tag = self._get_or_create_tag('BANK', raw)
                    self._record_value(tag, raw)
                    return tag

            if CTX_BANK.search(pre+post):
                tag = self._get_or_create_tag('BANK', raw)
                self._record_value(tag, raw)
                return tag
            if CTX_OP.search(pre+post):
                tag = self._get_or_create_tag('ID_CARD', raw)
                self._record_value(tag, raw)
                return tag

            return raw
        text = ACCT_RE.sub(acct_like, text)

        # DŮLEŽITÉ: IČO a DIČ PŘED IDCARD_RE!
        # Jinak "CZ28547896" se detekuje jako ID_CARD místo DIČ

        # IČO (Identifikační číslo organizace)
        def ico_repl(m):
            full_match = m.group(0)
            ico_num = m.group(1)
            tag = self._get_or_create_tag('ICO', ico_num)
            self._record_value(tag, ico_num)
            # Replace just the number, keep the label
            return full_match.replace(ico_num, tag)
        text = ICO_RE.sub(ico_repl, text)

        # DIČ (Daňové identifikační číslo)
        def dic_repl(m):
            full_match = m.group(0)
            dic_num = m.group(1)
            tag = self._get_or_create_tag('DIC', dic_num)
            self._record_value(tag, dic_num)
            # Replace just the number, keep the label
            return full_match.replace(dic_num, tag)
        text = DIC_RE.sub(dic_repl, text)

        # KRITICKÁ OPRAVA: Platební karty (PŘED IBAN)
        def card_repl(m):
            card_num = m.group(1)
            tag = self._get_or_create_tag('CARD', card_num)
            self._record_value(tag, card_num)
            return m.group(0).replace(card_num, tag)
        text = CARD_RE.sub(card_repl, text)

        # GDPR: IBAN (mezinárodní bankovní účet)
        def iban_repl(m):
            iban_num = m.group(1)
            # Normalizuj IBAN (odstraň mezery pro ukládání)
            iban_normalized = iban_num.replace(' ', '')
            tag = self._get_or_create_tag('IBAN', iban_normalized)
            self._record_value(tag, iban_normalized)
            return m.group(0).replace(iban_num, tag)
        text = IBAN_RE.sub(iban_repl, text)

        # GDPR: BIC/SWIFT (identifikátor banky) - s kontrolou kontextu
        # KRITICKÁ OPRAVA: "SYNERGIE" není BIC, je to název projektu
        BIC_BLACKLIST = {'synergie', 'project', 'projekt', 'alliance', 'aliance'}
        def bic_repl(m):
            v = m.group(1)  # BIC_RE má capturing group
            v_lower = v.lower()

            # Blacklist běžných slov (projektové názvy atd.)
            if v_lower in BIC_BLACKLIST:
                return m.group(0)  # Neanonymizuj

            # Kontext check: BIC by měl být poblíž "BIC", "SWIFT", "kód banky" atd.
            s, e = m.span()
            pre = text[max(0, s-50):s]
            post = text[e:e+50]

            if re.search(r'\b(BIC|SWIFT|kód\s+banky|bankovní\s+kód)\b', pre+post, re.IGNORECASE):
                tag = self._get_or_create_tag('BIC', v)
                self._record_value(tag, v)
                return tag

            # Pokud není bankovní kontext, neanonymizuj
            return m.group(0)

        text = BIC_RE.sub(bic_repl, text)

        def birth_or_id_repl(m):
            v = m.group(0)
            s, e = m.span()
            pre = text[max(0, s-40):s]
            post = text[e:e+40]

            # KRITICKÁ POLITIKA: Shape má přednost před labelem!
            # Pokud má tvar RČ (6 číslic / 3-4 číslice) → VŽDY [[BIRTH_ID_*]]
            # I když je kontext "Číslo OP:", fyzicky je to rodné číslo
            # Normalizuj číslo (odstraň mezery kolem lomítka)
            v_normalized = re.sub(r'\s*/\s*', '/', v)
            if re.match(r'^\d{6}/\d{3,4}$', v_normalized):
                tag = self._get_or_create_tag('BIRTH_ID', v)
                self._record_value(tag, v)
                return tag

            # DŮLEŽITÉ: Kontroluj CTX_BIRTH PŘED CTX_OP!
            # "Rodné číslo: 925315/6847 Číslo OP: 123" by jinak bylo ID_CARD kvůli "OP"

            # 1. Kontrola kontextu "r.č." nebo "(r.č." - pokud je tam, je to BIRTH_ID
            if re.search(r'[\(\s]r\.?\s*č\.?\s*[:\)]?\s*$', pre, re.IGNORECASE):
                tag = self._get_or_create_tag('BIRTH_ID', v)
            # 2. Kontrola "Rodné číslo:" PŘED číslem
            elif CTX_BIRTH.search(pre):
                tag = self._get_or_create_tag('BIRTH_ID', v)
            # 3. Teprve pak kontroluj OP kontext
            elif CTX_OP.search(pre+post):
                tag = self._get_or_create_tag('ID_CARD', v)
            # 4. Default je BIRTH_ID (formát 6/3-4 je nejčastěji RČ)
            else:
                tag = self._get_or_create_tag('BIRTH_ID', v)

            self._record_value(tag, v)
            return tag
        text = BIRTHID_RE.sub(birth_or_id_repl, text)

        def id_repl(m):
            v = m.group(0)
            s, e = m.span()

            # KRITICKÁ POLITIKA: Shape má přednost před labelem!
            # Pokud má tvar RČ (6 číslic / 3-4 číslice) → VŽDY [[BIRTH_ID_*]]
            # I když je kontext "Číslo OP:", fyzicky je to rodné číslo
            if re.match(r'^\d{6}/\d{3,4}$', v):
                tag = self._get_or_create_tag('BIRTH_ID', v)
                self._record_value(tag, v)
                return tag

            # Jinak je to ID_CARD (občanský průkaz)
            tag = self._get_or_create_tag('ID_CARD', v)
            self._record_value(tag, v)
            return tag
        text = IDCARD_RE.sub(id_repl, text)

        # KRITICKÁ OPRAVA: IP adresy
        def ip_repl(m):
            ip_addr = m.group(1)
            tag = self._get_or_create_tag('IP', ip_addr)
            self._record_value(tag, ip_addr)
            return m.group(0).replace(ip_addr, tag)
        text = IP_RE.sub(ip_repl, text)

        # KRITICKÁ OPRAVA: Hesla (NIKDY neukládat hodnotu do mapy!)
        def password_repl(m):
            password_value = m.group(1)
            # Vytvoř tag ale NEUKLÁDEJ hodnotu (bezpečnost!)
            self.counter['PASSWORD'] += 1
            tag = f'[[PASSWORD_{self.counter["PASSWORD"]}]]'
            # Zaznamenej pouze placeholder, ne skutečné heslo
            self.tag_map[tag] = ['********']
            return m.group(0).replace(password_value, tag)
        text = PASSWORD_RE.sub(password_repl, text)

        # KRITICKÁ OPRAVA: API klíče a tajemství (NIKDY neukládat hodnotu!)
        def api_key_repl(m):
            api_key_value = m.group(1)
            # Vytvoř tag ale NEUKLÁDEJ hodnotu (bezpečnost!)
            self.counter['API_KEY'] += 1
            tag = f'[[API_KEY_{self.counter["API_KEY"]}]]'
            # Zaznamenej pouze placeholder, ne skutečný klíč
            self.tag_map[tag] = ['********']
            return m.group(0).replace(api_key_value, tag)
        text = API_KEY_RE.sub(api_key_repl, text)

        # KRITICKÁ OPRAVA: Usernames a loginy
        def username_repl(m):
            username = m.group(1)
            tag = self._get_or_create_tag('USERNAME', username)
            self._record_value(tag, username)
            return m.group(0).replace(username, tag)
        text = USERNAME_RE.sub(username_repl, text)

        # KRITICKÁ OPRAVA: Čísla pojištěnců
        def insurance_repl(m):
            insurance_num = m.group(1)
            tag = self._get_or_create_tag('INSURANCE_ID', insurance_num)
            self._record_value(tag, insurance_num)
            return m.group(0).replace(insurance_num, tag)
        text = INSURANCE_ID_RE.sub(insurance_repl, text)

        # KRITICKÁ OPRAVA: RFID karty
        def rfid_repl(m):
            rfid_num = m.group(1)
            tag = self._get_or_create_tag('RFID', rfid_num)
            self._record_value(tag, rfid_num)
            return m.group(0).replace(rfid_num, tag)
        text = RFID_RE.sub(rfid_repl, text)

        # KRITICKÁ OPRAVA: Řidičské průkazy
        def driver_license_repl(m):
            license_num = m.group(1).strip()
            tag = self._get_or_create_tag('DRIVER_LICENSE', license_num)
            self._record_value(tag, license_num)
            return m.group(0).replace(m.group(1), tag)
        text = DRIVER_LICENSE_RE.sub(driver_license_repl, text)

        # Osobní číslo zaměstnance
        def emp_id_repl(m):
            full_match = m.group(0)
            emp_num = m.group(1)
            tag = self._get_or_create_tag('EMP_ID', emp_num)
            self._record_value(tag, emp_num)
            # Replace just the number, keep the label
            return full_match.replace(emp_num, tag)
        text = EMP_ID_RE.sub(emp_id_repl, text)

        return text

    def end_scan_for_leaks(self, text: str) -> list:
        """
        KRITICKÁ FUNKCE: Kontrola zbylých leaků po anonymizaci
        Vrací seznam nalezených leaků pro audit
        """
        leaks = []

        # Kontrola IBANů (včetně nalepených za ]])
        iban_pattern = re.compile(r'(?:\]\])?([A-Z]{2}\s?\d{2}(?:\s?\d{4}){3,7})(?!\]\])', re.IGNORECASE)
        for m in iban_pattern.finditer(text):
            if '[[IBAN_' not in text[max(0, m.start()-20):m.start()]:
                leaks.append(f"IBAN leak: {m.group(1)} at position {m.start()}")

        # Kontrola platebních karet
        card_pattern = re.compile(r'(?:\]\])?(\d{4}[\s\-]\d{4}[\s\-]\d{4}[\s\-]\d{4,7})(?!\]\])')
        for m in card_pattern.finditer(text):
            if '[[CARD_' not in text[max(0, m.start()-20):m.start()]:
                leaks.append(f"CARD leak: {m.group(1)[:7]}... at position {m.start()}")

        # Kontrola IP adres
        ip_pattern = re.compile(r'(?:\]\])?(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})(?!\]\])')
        for m in ip_pattern.finditer(text):
            if '[[IP_' not in text[max(0, m.start()-20):m.start()]:
                leaks.append(f"IP leak: {m.group(1)} at position {m.start()}")

        # Kontrola hesel
        password_pattern = re.compile(r'(?:Initial\s+password|Password|Heslo)\s*:?\s*([^\[\s]\S+)', re.IGNORECASE)
        for m in password_pattern.finditer(text):
            if '[[PASSWORD_' not in text[max(0, m.start()-20):m.start()]:
                leaks.append(f"PASSWORD leak at position {m.start()}")

        # Kontrola API klíčů
        api_pattern = re.compile(r'(?:AWS\s+Access\s+Key|AWS\s+Secret|API\s+Key|Stripe|SendGrid)\s*:?\s*([^\[\s][A-Za-z0-9+/=]{20,})', re.IGNORECASE)
        for m in api_pattern.finditer(text):
            if '[[API_KEY_' not in text[max(0, m.start()-20):m.start()]:
                leaks.append(f"API_KEY leak at position {m.start()}")

        # Kontrola usernames
        username_pattern = re.compile(r'(?:Login|Username|User)\s*:?\s*([^\[\s][A-Za-z0-9._\-@]+)', re.IGNORECASE)
        for m in username_pattern.finditer(text):
            if '[[USERNAME_' not in text[max(0, m.start()-20):m.start()]:
                leaks.append(f"USERNAME leak: {m.group(1)} at position {m.start()}")

        return leaks

    def post_merge_person_tags(self, doc: Document):
        key_to_tags = defaultdict(set)
        for tag, vals in list(self.tag_map.items()):
            if not tag.startswith('[[PERSON_'):
                continue
            for v in vals:
                m = PAIR_RE.search(v)
                if not m:
                    continue
                f_nom = infer_first_name_nominative(m.group(1), m.group(2)) or m.group(1)
                l_nom = infer_surname_nominative(m.group(2))
                key = (normalize_for_matching(f_nom), normalize_for_matching(l_nom))
                key_to_tags[key].add(tag)

        redirect = {}
        for key, tags in key_to_tags.items():
            if len(tags) <= 1:
                continue
            canon = sorted(tags)[0]
            for t in tags:
                if t != canon:
                    redirect[t] = canon

        if redirect:
            for p in iter_paragraphs(doc):
                txt = get_text(p)
                new = txt
                for src, dst in redirect.items():
                    new = new.replace(src, dst)
                if new != txt:
                    set_text(p, new)

            for src, dst in redirect.items():
                if src in self.tag_map:
                    for v in self.tag_map[src]:
                        if v not in self.tag_map[dst]:
                            self.tag_map[dst].append(v)
                    del self.tag_map[src]

    def anonymize_docx(self, input_path: str, output_path: str, json_map: str, txt_map: str):
        doc = Document(input_path)
        pieces = []
        for p in iter_paragraphs(doc):
            pieces.append(clean_invisibles(get_text(p)))
        self.source_text = '\n'.join(pieces)

        # KRITICKÁ OPRAVA: Před detekcí osob DOČASNĚ nahradit e-maily placeholdery
        # Jinak se jména v e-mailech (např. "martina.horáková@example.com") detekují jako osoby
        text_for_person_detection = EMAIL_RE.sub('__EMAIL_PLACEHOLDER__', self.source_text)

        self._extract_persons_to_index(text_for_person_detection)

        for p in iter_paragraphs(doc):
            raw = get_text(p)
            if not raw.strip():
                continue
            txt = clean_invisibles(raw)
            # DŮLEŽITÉ: Adresy MUSÍ být anonymizovány PŘED osobami!
            # Jinak "Novákova 45" končí jako "[[PERSON]] 45"
            txt = self.anonymize_entities(txt)  # Adresy, IČO, DIČ, telefony, emaily - PRVNÍ!
            txt = self._apply_known_people(txt)  # Potom známé osoby
            txt = self._replace_remaining_people(txt)  # Nakonec zbylé osoby
            if txt != raw:
                set_text(p, txt)

        self.post_merge_person_tags(doc)

        # KRITICKÁ KONTROLA: End-scan pro detekci zbylých leaků
        final_text_pieces = []
        for p in iter_paragraphs(doc):
            final_text_pieces.append(get_text(p))
        final_text = '\n'.join(final_text_pieces)

        leaks = self.end_scan_for_leaks(final_text)
        if leaks:
            print("\n⚠️  VAROVÁNÍ: End-scan našel potenciální leaky:")
            for leak in leaks:
                print(f"   - {leak}")
            print("⚠️  Doporučuji zkontrolovat výstupní dokument!\n")

        # Post-processing: Normalizace mezer kolem tagů (kosmetika pro enterprise reports)
        # Zajistí správné mezery: "Tel.:[[PHONE]]" → "Tel.: [[PHONE]]", "[[EMAIL]],[[PHONE]]" → "[[EMAIL]], [[PHONE]]"
        for p in iter_paragraphs(doc):
            txt = get_text(p)
            if '[[' in txt:
                # Oprava: ":" následované tagem bez mezery → přidat mezeru
                txt = re.sub(r':(\[\[)', r': \1', txt)
                # Oprava: "." následované tagem bez mezery → přidat mezeru (tel.[[PHONE]])
                txt = re.sub(r'\.(\[\[)', r'. \1', txt)
                # Oprava: "," následované tagem bez mezery → přidat mezeru ([[EMAIL]],[[PHONE]])
                txt = re.sub(r',(\[\[)', r', \1', txt)
                # Oprava: více mezer kolem tagů → jedna mezera
                txt = re.sub(r'\s{2,}', ' ', txt)
                set_text(p, txt)

        doc.save(output_path)

        data = OrderedDict((tag, self.tag_map[tag]) for tag in sorted(self.tag_map.keys()))
        with open(json_map, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        with open(txt_map, 'w', encoding='utf-8') as f:
            sections = [
                ("OSOBY", "PERSON"),
                ("RODNÁ ČÍSLA", "BIRTH_ID"),
                ("IČO", "ICO"),
                ("DIČ", "DIC"),
                ("OSOBNÍ ČÍSLA ZAMĚSTNANCŮ", "EMP_ID"),
                ("BANKOVNÍ ÚČTY", "BANK"),
                ("IBAN", "IBAN"),
                ("BIC/SWIFT", "BIC"),
                ("PLATEBNÍ KARTY", "CARD"),
                ("TELEFONY", "PHONE"),
                ("EMAILY", "EMAIL"),
                ("OBČANSKÉ PRŮKAZY", "ID_CARD"),
                ("ŘIDIČSKÉ PRŮKAZY", "DRIVER_LICENSE"),
                ("POZNÁVACÍ ZNAČKY (SPZ/RZ)", "LICENSE_PLATE"),
                ("VIN (VOZIDLA)", "VIN"),
                ("ČÍSLA POJIŠTĚNCŮ", "INSURANCE_ID"),
                ("RFID KARTY", "RFID"),
                ("IP ADRESY", "IP"),
                ("USERNAMES/ÚČTY", "USERNAME"),
                ("HESLA", "PASSWORD"),
                ("API KLÍČE", "API_KEY"),
                ("ČÁSTKY", "AMOUNT"),
                ("DATA", "DATE"),
                ("ADRESY", "ADDRESS"),
                ("MÍSTA NAROZENÍ", "PLACE"),
            ]
            for title, pref in sections:
                items = []
                for tag, vals in sorted(self.tag_map.items()):
                    if tag.startswith(f'[[{pref}_'):
                        if pref == "PERSON" and len(vals) > 0:
                            # Pro PERSON: první hodnota je kanonická, zbytek jsou varianty
                            canonical = vals[0]
                            items.append(f"{tag}: {canonical}")
                            if len(vals) > 1:
                                variants = vals[1:]
                                # Přidej varianty s odsazením
                                for v in variants:
                                    items.append(f"  - {v}")
                        else:
                            # Pro ostatní kategorie: standardní formát
                            for v in vals:
                                items.append(f"{tag}: {v}")
                if items:
                    f.write(f"{title}\n{'-'*len(title)}\n")
                    f.write("\n".join(items) + "\n\n")

def main():
    import argparse
    ap = argparse.ArgumentParser(description="Anonymizace českých DOCX s JSON knihovnou jmen")
    ap.add_argument("docx_path", nargs='?', help="Cesta k .docx souboru")
    ap.add_argument("--names-json", default="cz_names.v1.json", help="Cesta k JSON knihovně jmen")
    args = ap.parse_args()

    try:
        if args.names_json != "cz_names.v1.json":
            global CZECH_FIRST_NAMES
            CZECH_FIRST_NAMES = load_names_library(args.names_json)

        path = Path(args.docx_path) if args.docx_path else Path(input("Přetáhni sem .docx soubor nebo napiš cestu: ").strip().strip('"'))
        if not path.exists():
            print("❌ Soubor nenalezen:", path)
            input("\nStiskni Enter pro ukončení...")
            return 2

        base = path.stem
        out_docx = path.parent / f"{base}_anon.docx"
        out_json = path.parent / f"{base}_map.json"
        out_txt  = path.parent / f"{base}_map.txt"

        # Kontrola, zda výstupní soubory nejsou otevřené
        # Pokud ano, vytvoř nový soubor s časovým razítkem
        files_locked = False
        for out_file in [out_docx, out_json, out_txt]:
            if out_file.exists():
                try:
                    # Pokus se otevřít soubor pro zápis (testuje, zda není zamčený)
                    with open(out_file, 'a'):
                        pass
                except PermissionError:
                    files_locked = True
                    break

        if files_locked:
            # Vytvoř nové názvy souborů s časovým razítkem
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            out_docx = path.parent / f"{base}_anon_{timestamp}.docx"
            out_json = path.parent / f"{base}_map_{timestamp}.json"
            out_txt  = path.parent / f"{base}_map_{timestamp}.txt"
            print(f"\n⚠️  Výstupní soubory jsou otevřené v jiné aplikaci!")
            print(f"   Vytvářím nové soubory s časovým razítkem: {timestamp}")
            print()

        print(f"\n🔍 Zpracovávám: {path.name}")
        a = Anonymizer(verbose=False)
        a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))

        print("\n✅ Výstupy:")
        print(f" - {out_docx}")
        print(f" - {out_json}")
        print(f" - {out_txt}")
        print(f"\n📊 Statistiky:")
        print(f" - Nalezeno osob: {len(a.canonical_persons)}")
        print(f" - Celkem tagů: {sum(a.counter.values())}")

        # Pauza na konci pouze pokud je interaktivní terminál
        if sys.stdin.isatty():
            input("\n✅ Hotovo! Stiskni Enter pro ukončení...")
        return 0

    except Exception as e:
        print(f"\n❌ CHYBA: {e}")
        print(f"\n📋 Detail chyby:")
        import traceback
        traceback.print_exc()
        # Vždy pauza při chybě, aby uživatel viděl co se stalo
        try:
            input("\n⚠️  Stiskni Enter pro ukončení...")
        except:
            # Pokud input() selže, aspoň čekej 10 sekund
            import time
            print("\n⚠️  Zavírám za 10 sekund...")
            time.sleep(10)
        return 1

if __name__ == "__main__":
    sys.exit(main())